#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build three professional university project proposals (DOCX) + a combined master
document for the RISE MSAI Machine Learning course, Riphah International University.

Sections per proposal (per the brief):
 Cover Page, Table of Contents, Abstract, Introduction, Problem Statement, Objectives,
 Scope, Literature Review, Proposed Methodology, Technologies & Tools, Dataset Information,
 Expected Outcomes, Innovation/Uniqueness, Feasibility Analysis, Hardware & Software
 Requirements, Future Enhancements, Timeline / Gantt Chart, Conclusion, References.
"""
import os, copy, datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

ASSETS = "assets"
OUT = "."
TODAY = datetime.date(2026, 5, 11).strftime("%d %B %Y")   # submission date

UNIV   = "Riphah International University"
DEPT   = "RISE MSAI"
SUBJECT = "Machine Learning"
SUBMITTED_TO = "Junaid Khan"
MEMBERS = ["Azhar Hussain", "Rahma", "Haris Javaid", "Syed Farakh Abbas"]

# ---------------------------------------------------------------------------
# Theme colours mirroring gen_images.py
# ---------------------------------------------------------------------------
THEMES = {
    1: dict(name="AI-Based Crowd Counting System",
            short="Crowd Counting System",
            dark="0B132B", mid="1C2541", accent="1F6FEB", accent2="00B4D8",
            light="5BC0BE", lighttint="E3F1F4", headtint="DCEBF6"),
    2: dict(name="AI-Based Handwritten Text Recognition System",
            short="Handwritten Text Recognition",
            dark="2B0B3F", mid="3D1A52", accent="7B2CBF", accent2="9D4EDD",
            light="C77DFF", lighttint="F0E6F7", headtint="EBDFF5"),
    3: dict(name="Sign Language Recognition System",
            short="Sign Language Recognition",
            dark="06281F", mid="0B3D2E", accent="13795B", accent2="2D9B5C",
            light="7DDF9E", lighttint="E4F4E9", headtint="DDEEE3"),
}

# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------
def _set(el, **kw):
    for k, v in kw.items():
        el.set(qn(k), v)

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    _set(shd, **{"w:val": "clear", "w:color": "auto", "w:fill": hexcolor})
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}"); _set(e, **{"w:w": str(val), "w:type": "dxa"}); m.append(e)
    tcPr.append(m)

def set_cell_vertical_alignment(cell, val="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    e = OxmlElement("w:vAlign"); _set(e, **{"w:val": val}); tcPr.append(e)

def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        _set(e, **{"w:val": "nil"})
        borders.append(e)
    tblPr.append(borders)

def table_borders(table, color="BFBFBF", sz=6):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        _set(e, **{"w:val": "single", "w:sz": str(sz), "w:space": "0", "w:color": color})
        borders.append(e)
    tblPr.append(borders)

def add_field(paragraph, instr):
    """Insert a Word field (e.g. PAGE, NUMPAGES, TOC)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); _set(fldChar1, **{"w:fldCharType": "begin"})
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = instr
    fldChar2 = OxmlElement("w:fldChar"); _set(fldChar2, **{"w:fldCharType": "separate"})
    t = OxmlElement("w:t"); t.text = ""
    fldChar3 = OxmlElement("w:fldChar"); _set(fldChar3, **{"w:fldCharType": "end"})
    r = run._r
    r.append(fldChar1); r.append(instrText); r.append(fldChar2); r.append(t); r.append(fldChar3)
    return run

def set_run(run, size=None, bold=None, italic=None, color=None, font="Calibri", small_caps=False):
    if font: run.font.name = font
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color is not None: run.font.color.rgb = RGBColor.from_string(color)
    if small_caps: run.font.small_caps = True
    return run

def para(doc_or_cell, text="", size=11, bold=False, italic=False, color="222222",
         align=None, space_before=0, space_after=6, font="Calibri", line=1.15,
         style=None):
    p = doc_or_cell.add_paragraph(style=style) if hasattr(doc_or_cell, "add_paragraph") else doc_or_cell.add_paragraph()
    if align is not None: p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    if line: pf.line_spacing = line
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color, font=font)
    return p

def add_horizontal_rule(doc, color="BFBFBF", size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    _set(bottom, **{"w:val": "single", "w:sz": str(size), "w:space": "1", "w:color": color})
    bd.append(bottom); pPr.append(bd)
    return p

def page_break(doc):
    doc.add_page_break()

def add_picture_centered(doc, path, width_in, caption=None, th=None, space_after=10):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(path, width=Inches(width_in))
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_before = Pt(0); c.paragraph_format.space_after = Pt(space_after)
        r = c.add_run(caption)
        set_run(r, size=9, italic=True, color=(th["mid"] if th else "555555"))

# ---------------------------------------------------------------------------
# Headings with theme styling + bookmarks for TOC
# ---------------------------------------------------------------------------
_bookmark_id = [1000]

def heading(doc, text, level, th, number=None):
    """level 1 = section heading (Heading 1, themed band); level 2 = subheading."""
    p = doc.add_paragraph()
    style_name = "Heading 1" if level == 1 else ("Heading 2" if level == 2 else "Heading 3")
    p.style = doc.styles[style_name]
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(16); pf.space_after = Pt(8); pf.keep_with_next = True
        # shaded band
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); _set(shd, **{"w:val": "clear", "w:color": "auto", "w:fill": th["headtint"]})
        pPr.append(shd)
        bd = OxmlElement("w:pBdr")
        left = OxmlElement("w:left"); _set(left, **{"w:val": "single", "w:sz": "24", "w:space": "6", "w:color": th["accent"]})
        bottom = OxmlElement("w:bottom"); _set(bottom, **{"w:val": "single", "w:sz": "6", "w:space": "2", "w:color": th["accent"]})
        bd.append(left); bd.append(bottom); pPr.append(bd)
        ind = OxmlElement("w:ind"); _set(ind, **{"w:left": "120"}); pPr.append(ind)
        txt = (f"{number}.  " if number else "") + text
        r = p.add_run(txt); set_run(r, size=15, bold=True, color=th["dark"], font="Calibri")
    elif level == 2:
        pf.space_before = Pt(10); pf.space_after = Pt(4); pf.keep_with_next = True
        r = p.add_run(text); set_run(r, size=12.5, bold=True, color=th["accent"], font="Calibri")
        pPr = p._p.get_or_add_pPr()
        bd = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom"); _set(bottom, **{"w:val": "single", "w:sz": "4", "w:space": "2", "w:color": th["light"]})
        bd.append(bottom); pPr.append(bd)
    else:
        pf.space_before = Pt(8); pf.space_after = Pt(3); pf.keep_with_next = True
        r = p.add_run(text); set_run(r, size=11.5, bold=True, italic=True, color=th["mid"], font="Calibri")
    return p

def bullet(doc, text, th, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead:
        r = p.add_run(bold_lead + ": "); set_run(r, size=10.5, bold=True, color=th["dark"])
        r2 = p.add_run(text); set_run(r2, size=10.5, color="222222")
    else:
        r = p.add_run(text); set_run(r, size=10.5, color="222222")
    return p

def numbered(doc, text, th, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.12
    if bold_lead:
        r = p.add_run(bold_lead + " — "); set_run(r, size=10.5, bold=True, color=th["dark"])
        r2 = p.add_run(text); set_run(r2, size=10.5, color="222222")
    else:
        r = p.add_run(text); set_run(r, size=10.5, color="222222")
    return p

def body(doc, text, th, size=10.7, justify=True, space_after=8, bold=False, italic=False, color="222222"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text); set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p

# ---------------------------------------------------------------------------
# Generic table builder
# ---------------------------------------------------------------------------
def make_table(doc, headers, rows, th, col_widths=None, header_fs=9.5, body_fs=9.2,
               caption=None, zebra=True, first_col_bold=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    table_borders(t, color="C9C9C9", sz=4)
    # header
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        shade_cell(c, th["accent"])
        set_cell_margins(c); set_cell_vertical_alignment(c, "center")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
        r = p.add_run(str(h)); set_run(r, size=header_fs, bold=True, color="FFFFFF")
    # body
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            c = cells[j]; c.text = ""
            set_cell_margins(c); set_cell_vertical_alignment(c, "center")
            if zebra and i % 2 == 1:
                shade_cell(c, th["lighttint"])
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(val))
            set_run(r, size=body_fs, color="222222",
                    bold=(first_col_bold and j == 0))
            if first_col_bold and j == 0:
                r.font.color.rgb = RGBColor.from_string(th["dark"])
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_before = Pt(3); c.paragraph_format.space_after = Pt(10)
        r = c.add_run(caption); set_run(r, size=9, italic=True, color=th["mid"])
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

# ---------------------------------------------------------------------------
# Document-level setup: styles, page size, headers/footers
# ---------------------------------------------------------------------------
def setup_document(doc):
    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.7)
    style.paragraph_format.space_after = Pt(6)
    # A4 page, 1" margins
    for section in doc.sections:
        section.page_width = Cm(21.0); section.page_height = Cm(29.7)
        section.left_margin = Cm(2.4); section.right_margin = Cm(2.4)
        section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.2)
        section.header_distance = Cm(1.1); section.footer_distance = Cm(1.1)

def add_header_footer(section, th, title_text, different_first=True):
    section.different_first_page_header_footer = different_first
    # ---- header ----
    hdr = section.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # left: title, right: dept (use a tab)
    tabs = hp.paragraph_format
    r1 = hp.add_run(title_text); set_run(r1, size=8.5, color=th["mid"], italic=True)
    r2 = hp.add_run("\t\t" + f"{UNIV} — {DEPT}"); set_run(r2, size=8.5, color=th["mid"], italic=True)
    # bottom border on header
    pPr = hp._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom"); _set(bottom, **{"w:val": "single", "w:sz": "4", "w:space": "2", "w:color": th["accent"]})
    bd.append(bottom); pPr.append(bd)
    # set tab stop at right margin
    _add_right_tab(hp, section)
    # ---- footer ----
    ftr = section.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.text = ""
    pPr = fp._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    top = OxmlElement("w:top"); _set(top, **{"w:val": "single", "w:sz": "4", "w:space": "2", "w:color": th["light"]})
    bd.append(top); pPr.append(bd)
    rL = fp.add_run(f"{SUBJECT} Project Proposal"); set_run(rL, size=8.5, color=th["mid"])
    fp.add_run("\t")
    rC = fp.add_run("RISE MSAI"); set_run(rC, size=8.5, color=th["mid"])
    fp.add_run("\t")
    rP = fp.add_run("Page "); set_run(rP, size=8.5, color=th["mid"])
    pr = add_field(fp, "PAGE"); set_run(pr, size=8.5, color=th["mid"])
    rOf = fp.add_run(" of "); set_run(rOf, size=8.5, color=th["mid"])
    nr = add_field(fp, "NUMPAGES"); set_run(nr, size=8.5, color=th["mid"])
    _add_center_and_right_tabs(fp, section)
    # first-page footer (cover): minimal
    if different_first:
        ff = section.first_page_footer
        ff.is_linked_to_previous = False
        fpp = ff.paragraphs[0]; fpp.text = ""; fpp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fpp.add_run(f"{UNIV}  |  Department of {DEPT}  |  {SUBJECT}")
        set_run(r, size=8.5, color=th["mid"], italic=True)
        fh = section.first_page_header
        fh.is_linked_to_previous = False
        fh.paragraphs[0].text = ""

def _add_right_tab(paragraph, section):
    usable = section.page_width - section.left_margin - section.right_margin
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    t = OxmlElement("w:tab"); _set(t, **{"w:val": "right", "w:pos": str(int(usable))})
    tabs.append(t); pPr.append(tabs)

def _add_center_and_right_tabs(paragraph, section):
    usable = section.page_width - section.left_margin - section.right_margin
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tc = OxmlElement("w:tab"); _set(tc, **{"w:val": "center", "w:pos": str(int(usable / 2))})
    tr = OxmlElement("w:tab"); _set(tr, **{"w:val": "right", "w:pos": str(int(usable))})
    tabs.append(tc); tabs.append(tr); pPr.append(tabs)

def enable_update_fields(doc):
    """Tell Word to update fields (TOC, NUMPAGES) when the document opens."""
    settings = doc.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        settings.insert(0, el)
    _set(el, **{"w:val": "true"})

# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------
def cover_page(doc, th, pid, master=False, subtitle=None):
    # top thin spacing
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)
    # banner image
    add_picture_centered(doc, f"{ASSETS}/cover_{pid}.png", 6.5, caption=None, th=th, space_after=6)
    # "PROJECT PROPOSAL" tag
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("FINAL PROJECT PROPOSAL" if not master else "COMPILED PROJECT PROPOSALS")
    set_run(r, size=12, bold=True, color=th["accent"], small_caps=True); r.font.spacing = Pt(2)
    # Title
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(th["name"] if not master else "Three Machine Learning Project Proposals")
    set_run(r, size=22, bold=True, color=th["dark"])
    if subtitle:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(subtitle); set_run(r, size=11.5, italic=True, color=th["mid"])
    if master:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("1. Crowd Counting   •   2. Handwritten Text Recognition   •   3. Sign Language Recognition")
        set_run(r, size=10.5, italic=True, color=th["mid"])
    add_horizontal_rule(doc, color=th["accent"], size=12)
    # University block
    for txt, sz, bold, col in [
        (UNIV, 16, True, th["dark"]),
        (f"Department of {DEPT}", 12.5, True, th["mid"]),
        (f"Course / Subject:  {SUBJECT}", 11.5, False, "333333"),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(txt); set_run(r, size=sz, bold=bold, color=col)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    # Group members table
    gt = doc.add_table(rows=1, cols=2); gt.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_borders(gt, color=th["light"], sz=6)
    hdr = gt.rows[0].cells
    for j, h in enumerate(["Submitted By — Group Members", "Submitted To"]):
        shade_cell(hdr[j], th["accent"]); set_cell_margins(hdr[j], 70, 70, 140, 140)
        set_cell_vertical_alignment(hdr[j], "center")
        pp = hdr[j].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(h); set_run(r, size=10.5, bold=True, color="FFFFFF")
    row = gt.add_row().cells
    # members cell
    set_cell_margins(row[0], 90, 90, 150, 150); set_cell_vertical_alignment(row[0], "center")
    mp = row[0].paragraphs[0]; mp.paragraph_format.space_after = Pt(2)
    for i, m in enumerate(MEMBERS):
        if i: mp = row[0].add_paragraph()
        mp.paragraph_format.space_after = Pt(2)
        rr = mp.add_run(f"{i+1}.  {m}"); set_run(rr, size=11, color="222222", bold=(i == 0))
        if i == 0:
            rr2 = mp.add_run("   (Group Lead)"); set_run(rr2, size=9.5, italic=True, color=th["mid"])
    # submitted-to cell
    set_cell_margins(row[1], 90, 90, 150, 150); set_cell_vertical_alignment(row[1], "center")
    sp2 = row[1].paragraphs[0]; sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp2.add_run(SUBMITTED_TO); set_run(r, size=12.5, bold=True, color=th["dark"])
    sp3 = row[1].add_paragraph(); sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp3.add_run("Course Instructor — Machine Learning"); set_run(r, size=9.5, italic=True, color=th["mid"])
    gt.columns[0].width = Inches(3.7); gt.columns[1].width = Inches(2.7)
    for rrow in gt.rows:
        rrow.cells[0].width = Inches(3.7); rrow.cells[1].width = Inches(2.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    # submission date
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Date of Submission:  "); set_run(r, size=11, bold=True, color=th["dark"])
    r = p.add_run(TODAY); set_run(r, size=11, color="333333")
    add_horizontal_rule(doc, color=th["light"], size=8)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Academic Session 2025–2026   |   Confidential — For Academic Evaluation Only")
    set_run(r, size=8.5, italic=True, color=th["mid"])
    page_break(doc)

# ---------------------------------------------------------------------------
# Table of contents
# ---------------------------------------------------------------------------
def table_of_contents(doc, th, title="Table of Contents", levels="1-3"):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); _set(shd, **{"w:val": "clear", "w:color": "auto", "w:fill": th["headtint"]})
    pPr.append(shd)
    bd = OxmlElement("w:pBdr")
    left = OxmlElement("w:left"); _set(left, **{"w:val": "single", "w:sz": "24", "w:space": "6", "w:color": th["accent"]})
    bottom = OxmlElement("w:bottom"); _set(bottom, **{"w:val": "single", "w:sz": "6", "w:space": "2", "w:color": th["accent"]})
    bd.append(left); bd.append(bottom); pPr.append(bd)
    r = p.add_run(title); set_run(r, size=16, bold=True, color=th["dark"])
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    rn = note.add_run("If the entries below do not display page numbers in your viewer, place the cursor in "
                      "the table and press Ctrl+A then F9 (Word) — “Update Field → Update entire table”.")
    set_run(rn, size=8.5, italic=True, color="888888")
    # TOC field — use w:fldSimple so Word/LibreOffice recognise it as a real index
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    _set(fld, **{"w:instr": f' TOC \\o "{levels}" \\h \\z \\u '})
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); _set(sz, **{"w:val": "22"}); rpr.append(sz)
    col = OxmlElement("w:color"); _set(col, **{"w:val": "888888"}); rpr.append(col)
    r.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = "Table of contents — update this field to populate it."
    r.append(t); fld.append(r)
    p._p.append(fld)
    page_break(doc)

# ===========================================================================
#                          PROPOSAL CONTENT
# ===========================================================================
def section(doc, th, num, title):
    heading(doc, title, 1, th, number=num)

def proposal_body(doc, pid):
    th = THEMES[pid]
    # decorative divider after each major break occasionally
    DIV = f"{ASSETS}/divider_{pid}.png"
    ARCH = f"{ASSETS}/arch_{pid}.png"
    FLOW = f"{ASSETS}/flow_{pid}.png"
    GANTT = f"{ASSETS}/gantt_{pid}.png"
    CHART = f"{ASSETS}/chart_{pid}.png"

    if pid == 1:
        _proposal_1(doc, th, DIV, ARCH, FLOW, GANTT, CHART)
    elif pid == 2:
        _proposal_2(doc, th, DIV, ARCH, FLOW, GANTT, CHART)
    else:
        _proposal_3(doc, th, DIV, ARCH, FLOW, GANTT, CHART)


# ---------------------------------------------------------------------------
def _common_tech_table(doc, th, rows):
    make_table(doc, ["Category", "Tools / Technologies", "Purpose in the Project"], rows, th,
               col_widths=[1.4, 2.5, 2.5], first_col_bold=True,
               caption="Table: Technology stack selected for cost-free, low-resource development.")

# ===========================================================================
# PROPOSAL 1 — AI-BASED CROWD COUNTING SYSTEM
# ===========================================================================
def _proposal_1(doc, th, DIV, ARCH, FLOW, GANTT, CHART):
    # 2. ABSTRACT
    section(doc, th, 2, "Abstract")
    body(doc, "Crowd management has become a critical concern for public safety at religious gatherings, "
              "transport hubs, stadiums, shopping malls and political rallies. Manual head-counting from "
              "closed-circuit television (CCTV) is slow, error-prone and impractical in dense scenes where "
              "individuals heavily occlude one another. This project proposes an AI-Based Crowd Counting System "
              "that estimates the number of people and the spatial density of a crowd directly from CCTV frames "
              "using a lightweight density-map regression network. The core model is a transfer-learning variant "
              "of CSRNet — a VGG-16 front-end (frozen, pretrained on ImageNet) followed by a small dilated "
              "convolutional back-end that outputs a per-pixel density map whose integral equals the estimated "
              "head count. Around this model the system adds a colour heatmap overlay, a threshold-based "
              "over-crowding alert engine and a real-time analytics dashboard built with Streamlit. The model "
              "is trained on publicly available datasets (ShanghaiTech Part A/B and the Mall dataset) using a "
              "free Google Colab GPU, making the entire pipeline reproducible by university students without "
              "high-end hardware. Expected performance is a Mean Absolute Error (MAE) in the region of 70–95 "
              "on ShanghaiTech Part A and below 15 on the sparser Mall dataset, with near real-time inference "
              "(≈ 8–15 frames per second) on a mid-range laptop GPU or Colab. The deliverable is a deployable "
              "prototype dashboard plus a documented, retrainable model.", th)
    add_picture_centered(doc, DIV, 6.5, th=th)
    body(doc, "Keywords:  Crowd counting, density estimation, CSRNet, dilated convolution, transfer learning, "
              "computer vision, CCTV analytics, public-safety AI, Streamlit dashboard.", th, italic=True, size=10)

    # 3. INTRODUCTION
    section(doc, th, 3, "Introduction")
    body(doc, "The volume of video data generated by surveillance infrastructure has grown faster than the human "
              "capacity to analyse it. In Pakistan alone, large-scale events — Friday congregations at major "
              "mosques, Eid gatherings, Muharram processions, cricket matches at Gaddafi Stadium, and political "
              "jalsas — routinely draw tens of thousands of people into confined spaces. When density exceeds a "
              "safe threshold (commonly cited as 4–5 persons per square metre), crowd-crush incidents become a "
              "real risk. Automated crowd counting converts raw camera feeds into an actionable number and a "
              "spatial density field that operators can monitor in real time.", th)
    body(doc, "Early approaches to crowd counting relied on detecting individual heads or bodies with object "
              "detectors. These methods break down in dense crowds because of severe occlusion and scale "
              "variation. The field therefore moved toward density-map regression: instead of detecting people, "
              "a convolutional neural network (CNN) learns to predict a continuous density surface; summing the "
              "surface over an image region yields the count. Multi-Column CNN (MCNN) and, more influentially, "
              "CSRNet demonstrated that a VGG-based front-end combined with dilated convolutions captures both "
              "fine head-scale detail and large receptive-field context, achieving strong accuracy on benchmark "
              "datasets. Our project adapts this well-established, citation-rich architecture into a practical, "
              "resource-light system with an operational dashboard layer.", th)
    body(doc, "Crucially, this proposal is scoped to be implementable by a four-member student team within a "
              "single semester. We deliberately freeze the heavy VGG-16 feature extractor and train only a small "
              "back-end, use down-sampled training crops, and rely on free GPU quotas (Google Colab / Kaggle). "
              "The novelty is therefore not a new state-of-the-art network but a well-engineered, end-to-end, "
              "deployable application that fuses an accepted research model with heatmaps, alerting and "
              "analytics.", th)

    # 4. PROBLEM STATEMENT
    section(doc, th, 4, "Problem Statement")
    body(doc, "Public venues are monitored by CCTV operators who must visually estimate crowd size and judge "
              "when an area is becoming dangerously congested. This manual process suffers from four concrete "
              "problems:", th)
    bullet(doc, "Human counting is inaccurate beyond a few dozen people and degrades further with occlusion, "
                "perspective distortion and poor lighting.", th, bold_lead="Inaccuracy")
    bullet(doc, "One operator cannot continuously watch many feeds; fatigue and attention lapses cause missed "
                "build-ups.", th, bold_lead="Scalability")
    bullet(doc, "By the time over-crowding is noticed manually, the safe window for crowd dispersal may already "
                "have closed.", th, bold_lead="Latency")
    bullet(doc, "There is rarely a quantitative log of crowd levels over time for after-action review or capacity "
                "planning.", th, bold_lead="No analytics")
    body(doc, "Formally: given a CCTV image (or video stream) I, estimate (a) the total head count N, (b) a "
              "spatial density map D such that ∑D ≈ N, and (c) a binary over-crowding flag when local density "
              "or total count exceeds configurable thresholds — all in near real time and on commodity hardware. "
              "Existing research models solve (a) and (b) but are rarely packaged as a usable monitoring tool; "
              "this project closes that gap.", th)

    # 5. OBJECTIVES
    section(doc, th, 5, "Objectives")
    body(doc, "The project pursues the following measurable objectives:", th)
    for o in [
        ("Primary", "Design and implement a density-map regression model (CSRNet-style, transfer learning) "
                    "that estimates crowd count from a single CCTV frame with MAE ≤ 95 on ShanghaiTech Part A "
                    "and ≤ 15 on the Mall dataset."),
        ("Secondary", "Generate and overlay a colour density heatmap that visually highlights the most "
                      "congested regions of the scene."),
        ("Secondary", "Build a configurable over-crowding alert engine (per-region and global thresholds) that "
                      "raises a visual/audible warning and logs the event."),
        ("Secondary", "Develop a real-time analytics dashboard (Streamlit/Flask) showing live count, density "
                      "heatmap, time-series trend and alert history."),
        ("Secondary", "Achieve near real-time throughput (≥ 8 FPS on Colab GPU / mid-range laptop GPU) through "
                      "input down-scaling and a lightweight back-end."),
        ("Learning", "Demonstrate competence in CNNs, transfer learning, density-map ground-truth generation, "
                     "model evaluation (MAE/MSE) and ML application deployment."),
    ]:
        bullet(doc, o[1], th, bold_lead=o[0] + " objective")

    # 6. SCOPE
    section(doc, th, 6, "Scope of the Project")
    heading(doc, "In Scope", 2, th)
    for s in [
        "Image- and video-file based crowd counting (offline and near real-time playback).",
        "Single-camera, roughly fixed-viewpoint indoor/outdoor scenes similar to the training datasets.",
        "Density heatmap visualisation and PNG/CSV export of results.",
        "Threshold-based over-crowding alerts with an event log.",
        "A web dashboard with live count, heatmap, time-series chart and alert history.",
        "Model training/fine-tuning on ShanghaiTech and Mall datasets using free cloud GPUs.",
    ]:
        bullet(doc, s, th)
    heading(doc, "Out of Scope (Limitations)", 2, th)
    for s in [
        "Multi-camera 3-D person re-identification and tracking of specific individuals.",
        "Face recognition or any biometric identification (excluded for privacy and ethical reasons).",
        "Guaranteed accuracy on extreme scenes far outside the training distribution (e.g. aerial drone shots).",
        "Production-grade deployment, 24/7 uptime, or integration with proprietary VMS/CCTV servers.",
        "Edge-device (Raspberry Pi) deployment — feasible as a future enhancement but not a deliverable.",
    ]:
        bullet(doc, s, th)

    # 7. LITERATURE REVIEW
    section(doc, th, 7, "Literature Review")
    body(doc, "Crowd counting has evolved through three broad paradigms — detection-based, regression-based "
              "(global count regression), and density-map estimation — with the latter now dominant. The "
              "following review summarises the most relevant works and benchmark resources.", th)
    heading(doc, "7.1  Foundational and Recent Works", 2, th)
    body(doc, "Zhang et al. (2016) introduced the Multi-Column CNN (MCNN), using three CNN columns with "
              "different filter sizes to handle head-scale variation, and released the ShanghaiTech dataset that "
              "is still a primary benchmark. Li et al. (2018) proposed CSRNet — a VGG-16 front-end followed by "
              "dilated convolutions in the back-end — which removed the redundant multi-column design, enlarged "
              "the receptive field without losing resolution, and substantially reduced MAE; CSRNet is the "
              "architectural basis for this project. Subsequent work pushed accuracy further: Liu et al. (2019) "
              "Context-Aware Network (CAN) adaptively fuses multi-scale contextual features; Ma et al. (2019) "
              "Bayesian Loss reframes supervision around point annotations rather than blurred density maps; and "
              "Wang et al. (2020) DM-Count uses optimal-transport-based loss for sharper density estimates. "
              "Transformer-based counters (e.g. TransCrowd, 2021) and crowd-localisation methods continue to "
              "appear, but they are heavier to train and unnecessary for our resource-constrained scope.", th)
    body(doc, "For the application layer, OpenCV provides classical image I/O and pre-processing, while "
              "lightweight deployment tools (Streamlit, Flask, Gradio) make it straightforward to wrap a trained "
              "model in a real-time dashboard. The YOLO family (v5–v8) can serve as an optional sparse-scene "
              "person detector for cross-checking counts in low-density frames.", th)
    heading(doc, "7.2  Comparison of Crowd-Counting Approaches", 2, th)
    make_table(doc,
        ["Method (Year)", "Core Idea", "Strength", "Limitation / Cost", "Reported MAE (SH-A)"],
        [
            ["Head detection + counting", "Detect each person/head", "Interpretable; works when sparse",
             "Fails under heavy occlusion; slow", "≈ 90–110"],
            ["MCNN (2016)", "3-column CNN, multi-scale", "Handles scale variation; small model",
             "Lower accuracy; blurry maps", "110.2"],
            ["CSRNet (2018) ★ used", "VGG-16 + dilated conv back-end", "High accuracy; single stream; transfer-friendly",
             "VGG front-end is large (frozen here)", "68.2"],
            ["CAN (2019)", "Context-aware multi-scale fusion", "Better on scale-varying scenes",
             "More complex training", "62.3"],
            ["Bayesian Loss (2019)", "Point-supervision loss", "No need for tuned Gaussian maps",
             "Loss implementation overhead", "62.8"],
            ["DM-Count (2020)", "Optimal-transport loss", "Sharper, well-calibrated maps",
             "Heavier loss computation", "59.7"],
            ["TransCrowd (2021)", "Vision-Transformer regression", "Strong global context",
             "Large model; needs more GPU/data", "≈ 66"],
        ], th, col_widths=[1.3, 1.5, 1.5, 1.6, 1.0], header_fs=8.8, body_fs=8.6,
        caption="Table 1.1: Representative crowd-counting methods. ★ = approach adopted (in a frozen-front-end, "
                "low-resource configuration). MAE values are as reported by the original papers on ShanghaiTech Part A.")
    heading(doc, "7.3  Benchmark Datasets", 2, th)
    make_table(doc, ["Dataset", "Scale / Content", "Annotation", "Why Useful Here"],
        [
            ["ShanghaiTech Part A", "482 images, very dense web crowds", "Head-point annotations",
             "Primary benchmark; stress-tests dense counting"],
            ["ShanghaiTech Part B", "716 images, street scenes (sparser)", "Head-point annotations",
             "Closer to CCTV viewpoints; faster to train"],
            ["Mall dataset", "2 000 frames, single fixed camera", "Head-point annotations",
             "Realistic CCTV-style; small enough for laptops"],
            ["UCF-QNRF / JHU-CROWD++", "Large, extremely dense (optional)", "Head-point annotations",
             "Optional extra robustness; large download"],
        ], th, col_widths=[1.5, 1.9, 1.4, 2.1], header_fs=8.8, body_fs=8.6,
        caption="Table 1.2: Public crowd-counting datasets considered for training and evaluation.")
    body(doc, "Research gap addressed:  while accuracy-oriented papers abound, comparatively little student-level "
              "work packages a proven counter into a usable, real-time monitoring dashboard with alerting and "
              "historical analytics on free hardware. This project targets exactly that engineering gap.", th, italic=True)

    # 8. METHODOLOGY
    section(doc, th, 8, "Proposed Methodology")
    body(doc, "The system follows a classic train-once / serve-many design. The methodology is divided into a "
              "ground-truth preparation stage, a model training stage, and an application/serving stage.", th)
    heading(doc, "8.1  System Architecture", 2, th)
    add_picture_centered(doc, ARCH, 6.6, caption="Figure 1.1: End-to-end system architecture — from CCTV feed "
                         "through the CSRNet density model to heatmap, alerting and the analytics dashboard.", th=th)
    heading(doc, "8.2  Step-by-Step Workflow", 2, th)
    for i, (lead, txt) in enumerate([
        ("Data acquisition", "Download ShanghaiTech (Part A & B) and the Mall dataset; organise images and "
            "head-point annotation files (.mat / .json)."),
        ("Ground-truth density maps", "Convert each head-point annotation into a density map by placing a "
            "normalised 2-D Gaussian kernel at every head location (geometry-adaptive σ for dense scenes, fixed "
            "σ for sparser ones). The integral of the map equals the true count."),
        ("Pre-processing & augmentation", "Resize/crop images (e.g. 384–512 px on the short side), normalise "
            "with ImageNet statistics, and augment with random crops, horizontal flips and mild brightness/"
            "contrast jitter; rescale density maps consistently."),
        ("Model construction", "Build CSRNet: load VGG-16 (first 10 conv layers, ImageNet-pretrained) as a "
            "frozen front-end; attach a back-end of six dilated 3×3 convolutions (dilation 2) and a 1×1 conv "
            "producing a single-channel density map at 1/8 resolution."),
        ("Training", "Optimise pixel-wise Euclidean (MSE) loss between predicted and ground-truth density maps "
            "using Adam/SGD with a small learning rate; train on a free Colab/Kaggle GPU with checkpointing; "
            "monitor validation MAE and MSE."),
        ("Evaluation", "Report MAE and RMSE on the held-out test split; visualise predicted vs. ground-truth "
            "density maps; compare against an MCNN baseline."),
        ("Count extraction & heatmap", "Sum the predicted density map for the count; upsample and colour-map "
            "(jet/viridis) it, then alpha-blend over the original frame for the heatmap overlay."),
        ("Alert engine", "Compare global count and per-cell density against user-set thresholds; if exceeded, "
            "raise a dashboard alert (colour change + sound), and append {timestamp, count, location} to a log."),
        ("Dashboard & integration", "Wrap everything in a Streamlit app: upload image/video, see live count, "
            "heatmap, a rolling time-series chart and an alert table; allow CSV/PNG export."),
        ("Testing & optimisation", "Measure FPS, tune input resolution, optionally add YOLO cross-check for "
            "sparse frames, and document results."),
    ]):
        numbered(doc, txt, th, bold_lead=lead)
    heading(doc, "8.3  Workflow Flowchart", 2, th)
    add_picture_centered(doc, FLOW, 4.3, caption="Figure 1.2: Project workflow with the model-quality decision "
                         "gate and retraining loop.", th=th)
    heading(doc, "8.4  Algorithms, Models & Frameworks", 2, th)
    for b in [
        ("CSRNet (primary)", "VGG-16 front-end (frozen) + dilated-convolution back-end → single-channel density map."),
        ("MCNN (baseline)", "Lightweight multi-column CNN used as a comparison baseline."),
        ("YOLOv8-n (optional)", "Tiny person detector to cross-validate counts in low-density frames."),
        ("Loss & metrics", "Pixel-wise MSE training loss; evaluation by MAE and RMSE on counts."),
        ("Frameworks", "PyTorch (or TensorFlow/Keras), OpenCV, NumPy, SciPy (Gaussian kernels), Matplotlib, "
                       "Streamlit/Flask."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])

    # 9. TECHNOLOGIES AND TOOLS
    section(doc, th, 9, "Technologies and Tools")
    _common_tech_table(doc, th, [
        ["Programming language", "Python 3.10+", "Primary language for data prep, modelling and the dashboard"],
        ["Deep-learning framework", "PyTorch (alt. TensorFlow / Keras)", "Building, training and serving CSRNet/MCNN"],
        ["Computer vision", "OpenCV, Pillow", "Frame extraction, resizing, colour-mapping, video I/O"],
        ["Numerical / scientific", "NumPy, SciPy", "Array ops; geometry-adaptive Gaussian density-map generation"],
        ["Pretrained models", "torchvision VGG-16 (ImageNet), Ultralytics YOLOv8-n", "Transfer-learning front-end; optional sparse-scene detector"],
        ["Visualisation", "Matplotlib, Plotly", "Density-map plots, time-series charts in the dashboard"],
        ["App / dashboard", "Streamlit (alt. Flask + Bootstrap)", "Real-time upload, count, heatmap, alerts, analytics UI"],
        ["Training environment", "Google Colab / Kaggle Notebooks (free GPU)", "GPU training without local high-end hardware"],
        ["Dev tools", "VS Code, Jupyter, Git/GitHub", "Coding, experimentation, version control, collaboration"],
        ["Dataset sources", "ShanghaiTech, Mall dataset, (opt.) UCF-QNRF", "Training, validation and benchmarking data"],
    ])

    # 10. DATASET INFORMATION
    section(doc, th, 10, "Dataset Information")
    heading(doc, "10.1  Candidate Datasets", 2, th)
    for b in [
        ("ShanghaiTech Part A", "≈ 482 highly congested images crawled from the web; head-point annotations; "
                                "the standard hard benchmark."),
        ("ShanghaiTech Part B", "≈ 716 images of Shanghai street scenes; sparser, more CCTV-like; lighter to train."),
        ("Mall dataset", "≈ 2 000 frames from a single fixed shopping-mall camera; ideal for a realistic, "
                         "lightweight prototype."),
        ("Optional: UCF-QNRF / JHU-CROWD++", "Very large, extremely dense datasets for an additional robustness "
                                             "test if time and bandwidth permit."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])
    heading(doc, "10.2  Data Pre-processing", 2, th)
    for s in [
        "Parse head-point annotations and build per-image density maps using normalised 2-D Gaussian kernels "
        "(geometry-adaptive σ for dense Part A; fixed σ ≈ 4 px for sparser Part B / Mall).",
        "Resize images so the short side is ≈ 384–512 px (keeping aspect ratio); scale density maps to the "
        "network output resolution (1/8) and rescale values so the integral is preserved.",
        "Normalise pixels with ImageNet mean/standard deviation to match the pretrained VGG-16 front-end.",
        "Clean corrupt/duplicate files; verify that ∑(density map) ≈ annotated count for every sample.",
    ]:
        bullet(doc, s, th)
    heading(doc, "10.3  Training / Validation / Testing Split", 2, th)
    make_table(doc, ["Split", "Proportion", "Purpose"],
        [["Training", "≈ 70 %", "Learn back-end weights; the official train sets are used where provided"],
         ["Validation", "≈ 10–15 %", "Hyper-parameter tuning, early stopping, model selection"],
         ["Testing", "≈ 15–20 %", "Final unbiased MAE/RMSE; official test split used when available"]],
        th, col_widths=[1.4, 1.4, 3.6], first_col_bold=True,
        caption="Table 1.3: Data partitioning. Official ShanghaiTech splits are respected; the Mall dataset is "
                "split chronologically to avoid leakage between adjacent frames.")
    heading(doc, "10.4  Data Augmentation", 2, th)
    body(doc, "Because dense crowd datasets are relatively small, augmentation is essential: random crops "
              "(e.g. 256×256 / quarter-image patches), random horizontal flips, mild brightness/contrast/colour "
              "jitter, and occasional Gaussian noise or blur to mimic low-quality CCTV. All geometric transforms "
              "are applied identically to the image and its density map. Synthetic crowd imagery (e.g. GCC-style "
              "renders) can optionally be added for pre-training.", th)

    # 11. EXPECTED OUTCOMES
    section(doc, th, 11, "Expected Outcomes")
    for b in [
        "A trained CSRNet-style density model with MAE ≈ 70–95 on ShanghaiTech Part A and ≤ 15 on the Mall "
        "dataset (competitive with classic baselines, well within student reach).",
        "Accurate per-frame head counts plus a colour density heatmap overlay highlighting congestion hotspots.",
        "A working over-crowding alert engine with configurable thresholds and a timestamped event log.",
        "A real-time Streamlit dashboard: live count, heatmap, rolling time-series trend, alert history, CSV/PNG export.",
        "Near real-time throughput (≈ 8–15 FPS on Colab/mid-range GPU).",
        "Complete documentation: code repository, training notebooks, evaluation report and user guide.",
    ]:
        bullet(doc, b, th)
    add_picture_centered(doc, CHART, 6.6, caption="Figure 1.3: (left) MAE of representative methods on "
                         "ShanghaiTech Part A with our expected operating point; (right) an example predicted "
                         "density map whose integral approximates the crowd count.", th=th)

    # 12. INNOVATION
    section(doc, th, 12, "Innovation / Uniqueness")
    for b in [
        ("Research-to-product packaging", "Turns a citation-rich research model (CSRNet) into a complete, "
            "operator-facing monitoring tool — count + heatmap + alerts + analytics — which most student projects omit."),
        ("Low-resource engineering", "Frozen VGG-16 front-end, down-scaled crops and free cloud GPUs make "
            "state-of-research ideas reproducible on ordinary laptops."),
        ("Configurable safety layer", "Per-region and global density thresholds let the same model serve very "
            "different venues (mall corridor vs. stadium gate)."),
        ("Privacy-by-design", "Counts density, never identities — no face recognition — addressing a real "
            "ethical concern in public surveillance."),
        ("Analytics & after-action review", "Time-series logging supports capacity planning and incident review, "
            "not just live monitoring."),
        ("Optional hybrid cross-check", "A tiny YOLOv8-n detector validates counts in sparse frames, improving "
            "trust in the output."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])

    # 13. FEASIBILITY
    section(doc, th, 13, "Feasibility Analysis")
    heading(doc, "13.1  Technical Feasibility", 2, th)
    body(doc, "CSRNet, MCNN and YOLO have open-source reference implementations; the datasets are public; and "
              "PyTorch/OpenCV/Streamlit are mature. Freezing the VGG-16 front-end reduces the trainable "
              "parameters to a few million, so a single free Colab GPU session (T4, ~16 GB) can fine-tune the "
              "back-end in a few hours. Inference on down-scaled frames runs comfortably in near real time.", th)
    heading(doc, "13.2  Cost Analysis", 2, th)
    make_table(doc, ["Item", "Option Chosen", "Cost (PKR)"],
        [["Compute for training", "Google Colab / Kaggle free GPU tier", "0"],
         ["Datasets", "ShanghaiTech, Mall (public, free)", "0"],
         ["Software & libraries", "Python, PyTorch, OpenCV, Streamlit (open-source)", "0"],
         ["Development machines", "Members’ existing laptops", "0 (already owned)"],
         ["Optional Colab Pro (faster GPU)", "Only if needed for final tuning", "≈ 2 800 / month (optional)"],
         ["Hosting the demo (optional)", "Streamlit Community Cloud / Hugging Face Spaces (free)", "0"],
         ["Estimated total", "Core project", "≈ 0 (PKR 0–3 000 if Colab Pro used)"]],
        th, col_widths=[2.3, 2.6, 1.5], first_col_bold=True,
        caption="Table 1.4: Indicative cost analysis — the project is essentially zero-budget.")
    heading(doc, "13.3  Hardware Feasibility & GPU / Resource Limitations", 2, th)
    for b in [
        "Training: a free Colab/Kaggle GPU (NVIDIA T4/P100) is sufficient; full UCF-QNRF training is avoided due "
        "to size — Mall and ShanghaiTech are the working sets.",
        "Local development/inference: any laptop with 8 GB+ RAM; a 4 GB+ GPU is helpful but the model also runs "
        "on CPU at reduced FPS for demos.",
        "Mitigations for limited resources: smaller input resolution, gradient accumulation, mixed-precision "
        "training, checkpoint-and-resume to survive Colab session limits, and caching pre-computed density maps.",
    ]:
        bullet(doc, b, th)
    heading(doc, "13.4  Practical Implementation Plan", 2, th)
    body(doc, "Weeks 1–4: literature, dataset download, density-map pipeline, MCNN baseline. Weeks 5–8: build "
              "and train CSRNet on Colab, iterate to target MAE. Weeks 9–12: heatmap module, alert engine, "
              "Streamlit dashboard. Weeks 13–16: integration, FPS optimisation, testing, evaluation report, "
              "documentation and final demo. Work is split as: data/ground-truth (member A), model training "
              "(members A & B), dashboard/heatmap (member C), alerts/analytics & documentation (member D), with "
              "weekly integration checkpoints.", th)

    # 14. HARDWARE & SOFTWARE REQUIREMENTS
    section(doc, th, 14, "Hardware and Software Requirements")
    make_table(doc, ["Requirement", "Minimum", "Recommended"],
        [["Processor", "Quad-core CPU (Intel i5 / Ryzen 5)", "8-core CPU"],
         ["RAM", "8 GB", "16 GB"],
         ["GPU (local, optional)", "Integrated / 2 GB GPU (CPU inference OK)", "NVIDIA GPU 4–6 GB+ (e.g. GTX 1650)"],
         ["GPU (training)", "Google Colab / Kaggle free T4", "Colab Pro / institutional GPU"],
         ["Storage", "20 GB free", "50 GB free (datasets + checkpoints)"],
         ["OS", "Windows 10 / Ubuntu 20.04 / macOS", "Ubuntu 22.04"],
         ["Software", "Python 3.10, PyTorch, OpenCV, Streamlit, Git", "+ CUDA toolkit (if local GPU), VS Code"],
         ["Connectivity", "Broadband for dataset/model downloads", "Stable connection for Colab sessions"]],
        th, col_widths=[1.7, 2.4, 2.3], first_col_bold=True,
        caption="Table 1.5: Hardware and software requirements.")

    # 15. FUTURE ENHANCEMENTS
    section(doc, th, 15, "Future Enhancements")
    for b in [
        "Multi-camera fusion and approximate scene-level (homography-based) density maps.",
        "Edge deployment: model pruning/quantisation (TensorRT, ONNX, TFLite) for Jetson Nano / Raspberry Pi.",
        "Temporal smoothing with a lightweight RNN/Kalman filter for steadier video counts.",
        "Predictive crowding: short-horizon forecasting of density trends to warn before thresholds are hit.",
        "Anomaly detection (sudden surges, counter-flows) layered on the density stream.",
        "Mobile/operator app with push notifications; integration with standard VMS via RTSP.",
        "Domain adaptation / few-shot fine-tuning for a specific venue with minimal new labels.",
    ]:
        bullet(doc, b, th)

    # 16. TIMELINE
    section(doc, th, 16, "Timeline / Gantt Chart")
    body(doc, "The project spans a ≈ 16-week semester, organised into a foundation phase, a modelling & "
              "development phase, and a finalisation phase. Weekly milestones are shown below.", th)
    add_picture_centered(doc, GANTT, 6.7, caption="Figure 1.4: Project Gantt chart with weekly milestone breakdown.", th=th)
    make_table(doc, ["Week(s)", "Milestone / Deliverable"],
        [["W1–W2", "Requirement analysis, literature review, finalise architecture and tools"],
         ["W2–W3", "Download & organise ShanghaiTech and Mall datasets; exploratory analysis"],
         ["W3–W4", "Build density-map ground-truth pipeline; data-loader and augmentation"],
         ["W5–W6", "Implement MCNN baseline; set up Colab training environment"],
         ["W6–W8", "Build and train CSRNet (frozen VGG-16 + dilated back-end); reach target MAE"],
         ["W9–W10", "Heatmap generation module + over-crowding alert engine with event logging"],
         ["W10–W12", "Streamlit analytics dashboard (live count, heatmap, time-series, alerts)"],
         ["W12–W13", "End-to-end integration; near-real-time optimisation (resolution, batching)"],
         ["W13–W14", "Testing, evaluation, error analysis, comparison tables and charts"],
         ["W14–W15", "Documentation, user guide and final report writing"],
         ["W16", "Final review, demonstration and submission"]],
        th, col_widths=[1.1, 5.3], first_col_bold=True,
        caption="Table 1.6: Weekly milestone breakdown.")

    # 17. CONCLUSION
    section(doc, th, 17, "Conclusion")
    body(doc, "This proposal presented an AI-Based Crowd Counting System that converts ordinary CCTV footage "
              "into real-time head counts, a density heatmap, automated over-crowding alerts and historical "
              "analytics. The technical core — a CSRNet-style density-map regressor with a frozen, pretrained "
              "VGG-16 front-end — is well established in the literature yet light enough to train on free cloud "
              "GPUs and run on commodity laptops. The project’s contribution is engineering and integration: "
              "wrapping a proven model in a privacy-respecting, operator-friendly dashboard. It is realistic for "
              "a four-member team to deliver within one semester, develops core machine-learning competencies, "
              "and produces a genuinely useful public-safety tool with a clear path to future enhancement.", th)

    # 18. REFERENCES
    section(doc, th, 18, "References")
    refs_1 = [
        'Y. Zhang, D. Zhou, S. Chen, S. Gao, and Y. Ma, "Single-image crowd counting via multi-column '
        'convolutional neural network," in Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition (CVPR), 2016, pp. 589–597.',
        'Y. Li, X. Zhang, and D. Chen, "CSRNet: Dilated convolutional neural networks for understanding the '
        'highly congested scenes," in Proc. IEEE/CVF CVPR, 2018, pp. 1091–1100. [Online]. Available: https://arxiv.org/abs/1802.10062',
        'W. Liu, M. Salzmann, and P. Fua, "Context-aware crowd counting," in Proc. IEEE/CVF CVPR, 2019, pp. 5099–5108.',
        'Z. Ma, X. Wei, X. Hong, and Y. Gong, "Bayesian loss for crowd count estimation with point supervision," '
        'in Proc. IEEE/CVF Int. Conf. Computer Vision (ICCV), 2019, pp. 6142–6151.',
        'B. Wang, H. Liu, D. Samaras, and M. Hoai, "Distribution matching for crowd counting (DM-Count)," in '
        'Advances in Neural Information Processing Systems (NeurIPS), vol. 33, 2020, pp. 1595–1607.',
        'D. Liang, X. Chen, W. Xu, Y. Zhou, and X. Bai, "TransCrowd: Weakly-supervised crowd counting with '
        'transformers," Science China Information Sciences, vol. 65, no. 6, 2022.',
        'V. A. Sindagi and V. M. Patel, "A survey of recent advances in CNN-based single image crowd counting '
        'and density estimation," Pattern Recognition Letters, vol. 107, pp. 3–16, 2018.',
        'C. Wang, H. Zhang, L. Yang, S. Liu, and X. Cao, "Deep people counting in extremely dense crowds," in '
        'Proc. ACM Int. Conf. Multimedia (ACM MM), 2015, pp. 1299–1302.  (ShanghaiTech / dense-crowd context.)',
        'G. Jocher et al., "Ultralytics YOLOv8," 2023. [Online]. Available: https://github.com/ultralytics/ultralytics',
        'A. Paszke et al., "PyTorch: An imperative style, high-performance deep learning library," in NeurIPS, 2019, pp. 8024–8035.',
        'G. Bradski, "The OpenCV Library," Dr. Dobb’s Journal of Software Tools, 2000.',
        'Papers With Code, "ShanghaiTech crowd-counting dataset." [Online]. Available: https://paperswithcode.com/dataset/shanghaitech',
        'C. C. Loy, S. Gong, and T. Xiang, "From semi-supervised to transfer counting of crowds (Mall dataset)," '
        'in Proc. IEEE ICCV, 2013, pp. 2256–2263.',
        'Streamlit Inc., "Streamlit — A faster way to build and share data apps," documentation, 2024. [Online]. Available: https://docs.streamlit.io',
    ]
    _refs(doc, th, refs_1)


# ===========================================================================
# PROPOSAL 2 — AI-BASED HANDWRITTEN TEXT RECOGNITION
# ===========================================================================
def _proposal_2(doc, th, DIV, ARCH, FLOW, GANTT, CHART):
    section(doc, th, 2, "Abstract")
    body(doc, "Vast amounts of human knowledge — lecture notes, examination scripts, historical manuscripts, "
              "medical records, official registers, filled forms — still exist only as handwriting on paper. "
              "Converting this content into editable, searchable digital text unlocks indexing, translation, "
              "accessibility and long-term preservation. This project proposes an AI-Based Handwritten Text "
              "Recognition (HTR) system that takes an image, scanned page or PDF of handwritten text and outputs "
              "editable text, exportable as TXT, DOCX or a searchable PDF. The recognition engine is a "
              "Convolutional Recurrent Neural Network (CRNN): a CNN feature extractor (a compact 7-layer stack "
              "or a MobileNet-style backbone) feeds a bidirectional LSTM sequence encoder, which is decoded with "
              "Connectionist Temporal Classification (CTC) — the architecture introduced by Shi et al. for "
              "image-based sequence recognition. Around this core the system adds image pre-processing "
              "(deskewing, binarisation, line and word segmentation), a dictionary/edit-distance spell-correction "
              "post-processor (SymSpell), basic layout reconstruction, and an optional second model for Urdu/"
              "regional scripts. The system is trained on the IAM Handwriting Database (and synthetic data "
              "generated from handwriting fonts) using a free Google Colab GPU. Targeted performance is a "
              "word-level Character Error Rate (CER) of roughly 8–12 % on IAM lines without a language model and "
              "lower with post-processing — comfortably achievable by a student team — together with a usable "
              "desktop/web application for everyday note digitisation.", th)
    add_picture_centered(doc, DIV, 6.5, th=th)
    body(doc, "Keywords:  Handwritten text recognition, OCR, CRNN, CTC loss, BiLSTM, transfer learning, "
              "document digitisation, IAM database, multi-language OCR.", th, italic=True, size=10)

    section(doc, th, 3, "Introduction")
    body(doc, "Optical Character Recognition (OCR) for printed text is largely a solved problem, but handwritten "
              "text recognition remains challenging because handwriting is highly variable — slant, spacing, "
              "letter shapes, ligatures and pen pressure differ between writers and even within one writer’s "
              "page. Classical HTR pipelines first segmented a line into individual characters and classified "
              "each one, a brittle approach that fails when letters touch or overlap. The breakthrough came with "
              "segmentation-free sequence learning: an image of an entire word or text line is processed by a "
              "CNN to extract a feature sequence, a recurrent network models the left-to-right context, and a "
              "CTC layer aligns the variable-length prediction to the variable-length label without needing "
              "per-character segmentation. Shi, Bai and Yao’s CRNN (2015) established this end-to-end design and "
              "remains the practical workhorse of HTR; later work (Puigcerver 2017; Bluche & Messina 2017) "
              "refined the convolutional and recurrent blocks, and recent Transformer models such as TrOCR "
              "(2021) push accuracy further at much higher computational cost.", th)
    body(doc, "This project deliberately builds on the CRNN+CTC design because it is accurate, well documented, "
              "and trainable on a free cloud GPU. To make the system genuinely useful rather than a bare model, "
              "we add a complete document pipeline: pre-processing (deskew, binarise, denoise), line/word "
              "segmentation, the CRNN recogniser, a spell-correction post-processor, simple layout "
              "reconstruction, and export to editable formats. Multi-language support (an English model plus an "
              "optional Urdu/Arabic-script model) and a clean user interface complete the deliverable.", th)

    section(doc, th, 4, "Problem Statement")
    body(doc, "Organisations and individuals accumulate large volumes of handwritten material that cannot be "
              "searched, edited, translated or efficiently archived. Manual transcription is the current "
              "fallback, but it is:", th)
    bullet(doc, "Extremely time-consuming and expensive at scale (thousands of pages of notes, forms, registers).", th, bold_lead="Costly")
    bullet(doc, "Error-prone and inconsistent, especially for unfamiliar handwriting or domain jargon.", th, bold_lead="Inaccurate")
    bullet(doc, "A barrier to accessibility — handwritten notes are not screen-reader friendly for visually "
                "impaired users.", th, bold_lead="Inaccessible")
    bullet(doc, "An obstacle to preservation — paper degrades, and without digitisation the content is at risk.", th, bold_lead="Fragile")
    body(doc, "Formally: given an input image / scan / PDF page P containing handwritten text, produce an "
              "editable digital transcript T that closely matches the written content (low character and word "
              "error rates), preserving basic layout (lines, paragraphs), supporting more than one language/"
              "script, and exporting to TXT / DOCX / searchable PDF — all runnable on student hardware. Research "
              "models address the core recognition step; this project additionally delivers the surrounding, "
              "user-facing document workflow.", th)

    section(doc, th, 5, "Objectives")
    for o in [
        ("Primary", "Implement a CRNN + CTC handwriting recogniser achieving a line-level CER of roughly "
                    "8–12 % on the IAM test set without a language model (lower with post-processing)."),
        ("Secondary", "Build a robust pre-processing stage: orientation correction, binarisation, denoising, "
                      "and line/word segmentation for multi-line documents."),
        ("Secondary", "Add a dictionary + edit-distance spell-correction post-processor (SymSpell) to reduce "
                      "word-level errors."),
        ("Secondary", "Provide multi-language support — an English model plus an optional Urdu/Arabic-script "
                      "model — with a simple language selector."),
        ("Secondary", "Reconstruct basic layout (line order, paragraphs) and export results as TXT, DOCX and "
                      "searchable PDF."),
        ("Secondary", "Deliver a clean desktop/web UI for uploading images/PDFs and reviewing/editing the "
                      "recognised text, including a near-real-time webcam capture mode."),
        ("Learning", "Demonstrate competence in CNNs, RNN/LSTM sequence modelling, CTC loss, transfer learning, "
                     "data augmentation and OCR evaluation (CER/WER)."),
    ]:
        bullet(doc, o[1], th, bold_lead=o[0] + " objective")

    section(doc, th, 6, "Scope of the Project")
    heading(doc, "In Scope", 2, th)
    for s in [
        "Recognition of handwritten Latin-script (English) text from images, scans and PDF pages.",
        "Pre-processing: deskew, binarise, denoise, and line/word segmentation for multi-line documents.",
        "CRNN+CTC recognition with greedy and beam-search decoding; SymSpell post-correction.",
        "Optional second model for Urdu/Arabic-script words (proof-of-concept, smaller dataset).",
        "Basic layout reconstruction and export to TXT, DOCX and searchable PDF.",
        "A desktop/web application with file upload, webcam capture and an editable results view.",
    ]:
        bullet(doc, s, th)
    heading(doc, "Out of Scope (Limitations)", 2, th)
    for s in [
        "Full historical-manuscript transcription with complex multi-column layouts, marginalia and tables.",
        "Cursive recognition in arbitrary languages beyond the trained models; mathematical-formula recognition.",
        "Guaranteed accuracy on extremely messy, faded or low-resolution handwriting outside the training distribution.",
        "Large-scale Transformer models (e.g. full TrOCR fine-tuning) — noted as a future enhancement only.",
        "Production-grade hosting, user accounts and document-management features.",
    ]:
        bullet(doc, s, th)

    section(doc, th, 7, "Literature Review")
    body(doc, "Handwritten text recognition has progressed from per-character segmentation, through HMM- and "
              "MDLSTM-based line recognisers, to the now-standard CNN+RNN+CTC family and, most recently, "
              "Transformer-based encoder–decoder models. The works below frame our design choices.", th)
    heading(doc, "7.1  Key Works", 2, th)
    body(doc, "Graves et al. (2009) showed that multi-dimensional LSTMs with a CTC output layer could recognise "
              "unconstrained handwriting at the line level without explicit segmentation, winning ICDAR "
              "competitions. Shi, Bai and Yao (2015) introduced the CRNN — convolutional layers for feature "
              "extraction, bidirectional LSTMs for sequence modelling, and a CTC transcription layer — as a "
              "unified, end-to-end trainable network for image-based sequence recognition; this is the backbone "
              "of our system. Puigcerver (2017) demonstrated that carefully designed 1-D (rather than 2-D) "
              "recurrent layers on convolutional features are sufficient and faster, setting strong IAM "
              "baselines, while Bluche and Messina (2017) proposed gated CNN-BLSTM architectures for "
              "multilingual HTR. Graves (2012) and Graves et al. (2006) provide the theoretical basis of CTC. "
              "More recently, TrOCR (Li et al., 2021) couples a pretrained vision Transformer encoder with a "
              "pretrained text Transformer decoder, achieving very low error rates but requiring substantial "
              "compute and large pretraining corpora — beyond our scope, though an obvious future direction. "
              "Open-source engines such as Tesseract 4/5 (LSTM-based) and toolkits like PyLaia, Kraken and "
              "PaddleOCR offer reusable components and baselines. SymSpell (Garbe) supplies fast symmetric-"
              "delete spelling correction for the post-processing stage.", th)
    heading(doc, "7.2  Comparison of HTR Approaches", 2, th)
    make_table(doc,
        ["Approach (Year)", "Core Idea", "Strength", "Limitation / Cost", "IAM CER (approx.)"],
        [
            ["Char-segmentation + CNN", "Segment line→chars, classify each", "Simple components",
             "Brittle; touching letters break it", "high / unstable"],
            ["MDLSTM + CTC (2009)", "2-D LSTM over the image + CTC", "Segmentation-free; historic SOTA",
             "Slow; hard to train", "≈ 18 %"],
            ["CRNN (CNN+BLSTM+CTC) (2015) ★", "CNN features → BiLSTM → CTC", "Accurate, end-to-end, light, well-documented",
             "Needs decent data & augmentation", "≈ 10–12 %"],
            ["Puigcerver CNN-1D-LSTM (2017)", "Conv features + 1-D LSTMs", "Faster; strong IAM baseline",
             "Still needs careful tuning", "≈ 8 %"],
            ["Gated CNN-BLSTM (2017)", "Gated conv + BLSTM, multilingual", "Good cross-language results",
             "More complex architecture", "≈ 8 %"],
            ["TrOCR (2021)", "ViT encoder + Transformer decoder", "Very low error; no CTC needed",
             "Large model; heavy pretraining/compute", "≈ 3–4 %"],
            ["Tesseract 4/5 (LSTM)", "Open-source LSTM line recogniser", "Free, ready-to-use, many languages",
             "Weaker on free-form handwriting", "varies"],
        ], th, col_widths=[1.45, 1.55, 1.55, 1.55, 0.95], header_fs=8.6, body_fs=8.5,
        caption="Table 2.1: Representative HTR approaches. ★ = backbone adopted (CRNN + CTC, with a lightweight "
                "CNN front-end). CER figures are indicative line-level results reported in the literature on the IAM database.")
    heading(doc, "7.3  Datasets & Resources", 2, th)
    make_table(doc, ["Resource", "Content", "Use Here"],
        [["IAM Handwriting Database", "≈ 1 539 scanned forms, 13 353 labelled lines, 657 writers (English)",
          "Primary training/validation/testing data"],
         ["IAM-OnDB / CVL / RIMES", "Additional handwriting corpora (online strokes / multi-writer / French)",
          "Optional extra data and robustness checks"],
         ["Synthetic handwriting (TRDG / fonts)", "Text rendered with many handwriting fonts + distortions",
          "Pre-training and augmentation to enlarge the training set"],
         ["UrduDoc / UNHD / KHATT (Arabic-script)", "Urdu/Arabic handwritten text datasets",
          "Optional second-language proof-of-concept"],
         ["SymSpell dictionary, NLTK/WordFreq", "Word lists and frequencies", "Spell-correction post-processing"]],
        th, col_widths=[1.9, 2.6, 1.9], header_fs=8.8, body_fs=8.6,
        caption="Table 2.2: Datasets and resources considered.")
    body(doc, "Research gap addressed:  most academic HTR work optimises the recogniser in isolation on IAM "
              "lines. Far fewer student projects deliver the full document workflow — pre-processing, "
              "segmentation, post-correction, layout rebuild, multi-language, and export — wrapped in a usable "
              "app. This project targets that practical, end-to-end gap on free hardware.", th, italic=True)

    section(doc, th, 8, "Proposed Methodology")
    body(doc, "The methodology has three stages: a document pre-processing/segmentation stage, a CRNN+CTC "
              "recognition stage (trained once, served many times), and a post-processing/export/UI stage.", th)
    heading(doc, "8.1  System Architecture", 2, th)
    add_picture_centered(doc, ARCH, 6.6, caption="Figure 2.1: End-to-end architecture — image/scan/PDF → "
                         "pre-processing & segmentation → CRNN (CNN + BiLSTM + CTC) → spell-correction & layout "
                         "rebuild → editable TXT/DOCX/searchable-PDF output.", th=th)
    heading(doc, "8.2  Step-by-Step Workflow", 2, th)
    for lead, txt in [
        ("Input handling", "Accept JPG/PNG images, scanned pages, multi-page PDFs (rasterised), or near-real-"
            "time webcam frames."),
        ("Pre-processing", "Convert to grayscale; correct skew (Hough/projection-profile); binarise (Otsu / "
            "adaptive / Sauvola); remove noise and ruled lines; normalise contrast."),
        ("Segmentation", "Detect text regions, then segment into lines (horizontal projection / connected "
            "components) and, where needed, into words; normalise each line image height (e.g. 32–64 px)."),
        ("Recogniser construction", "Build the CRNN: a CNN feature extractor (compact 7-conv stack or "
            "MobileNet-style backbone, optionally ImageNet-initialised) → map-to-sequence → 2 stacked "
            "bidirectional LSTM layers → linear layer over the character set → CTC."),
        ("Training", "Train end-to-end with CTC loss on IAM (plus synthetic data) using Adam, learning-rate "
            "scheduling and gradient clipping on a free Colab/Kaggle GPU; checkpoint regularly; monitor "
            "validation CER/WER."),
        ("Decoding", "Use greedy (best-path) decoding for speed and CTC beam search (optionally with a small "
            "n-gram language model) for accuracy."),
        ("Post-processing", "Apply SymSpell dictionary correction; restore casing/punctuation heuristics; "
            "reassemble words → lines → paragraphs in reading order."),
        ("Multi-language", "Provide a language selector that swaps in an alternative model (e.g. an Urdu/Arabic-"
            "script CRNN) and an appropriate dictionary."),
        ("Export & UI", "Render the result in an editable text panel beside the source image; export to TXT, "
            "DOCX (python-docx) and searchable PDF (OCR text layer over the image)."),
        ("Testing & evaluation", "Measure CER/WER on held-out IAM data and on a small custom test set of team "
            "members’ handwriting; analyse error types; document results."),
    ]:
        numbered(doc, txt, th, bold_lead=lead)
    heading(doc, "8.3  Workflow Flowchart", 2, th)
    add_picture_centered(doc, FLOW, 4.3, caption="Figure 2.2: Project workflow with the model-quality decision "
                         "gate and refinement loop.", th=th)
    heading(doc, "8.4  Algorithms, Models & Frameworks", 2, th)
    for b in [
        ("CRNN (primary)", "CNN feature extractor + 2× BiLSTM + CTC transcription layer."),
        ("CTC loss / decoders", "Connectionist Temporal Classification training loss; greedy and beam-search decoding."),
        ("CNN backbone options", "Compact 7-layer conv stack (VGG-like) or MobileNetV3-small features (transfer learning)."),
        ("Pre-processing algorithms", "Otsu/Sauvola binarisation, projection-profile deskew & line segmentation, connected components."),
        ("Post-processing", "SymSpell symmetric-delete spelling correction; optional KenLM n-gram language model in beam search."),
        ("Baselines", "Tesseract 5 (LSTM) and a per-character CNN classifier for comparison."),
        ("Frameworks", "PyTorch / TensorFlow-Keras, OpenCV, NumPy, scikit-image, python-docx, pdf2image/PyMuPDF, Streamlit/PyQt."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])

    section(doc, th, 9, "Technologies and Tools")
    _common_tech_table(doc, th, [
        ["Programming language", "Python 3.10+", "Data prep, modelling, post-processing, application"],
        ["Deep-learning framework", "PyTorch (alt. TensorFlow / Keras)", "Building and training the CRNN; CTC loss"],
        ["Computer vision / imaging", "OpenCV, scikit-image, Pillow", "Binarisation, deskew, denoise, line/word segmentation"],
        ["Numerical", "NumPy", "Array operations, feature handling"],
        ["Pretrained backbones", "torchvision MobileNetV3 / ResNet18 (ImageNet)", "Optional transfer-learning CNN front-end"],
        ["OCR baselines / tools", "Tesseract 5, (ref.) PaddleOCR / PyLaia", "Comparison baselines and reusable components"],
        ["Post-processing / NLP", "SymSpell, NLTK / wordfreq, (opt.) KenLM", "Spell correction and language modelling"],
        ["Document I/O & export", "pdf2image / PyMuPDF, python-docx, reportlab", "Read PDFs; export TXT/DOCX/searchable PDF"],
        ["App / UI", "Streamlit (alt. PyQt / Flask)", "Upload, webcam capture, editable results view, export buttons"],
        ["Training environment", "Google Colab / Kaggle Notebooks (free GPU)", "GPU training without local high-end hardware"],
        ["Dev tools", "VS Code, Jupyter, Git/GitHub", "Coding, experiments, version control, collaboration"],
        ["Dataset sources", "IAM database; synthetic (TRDG); (opt.) Urdu/Arabic sets", "Training, validation, testing, multi-language PoC"],
    ])

    section(doc, th, 10, "Dataset Information")
    heading(doc, "10.1  Candidate Datasets", 2, th)
    for b in [
        ("IAM Handwriting Database", "≈ 1 539 scanned forms → 13 353 labelled text lines from 657 writers "
                                     "(English, modern Latin script); the standard HTR benchmark."),
        ("Synthetic handwriting data", "Text strings rendered with dozens of handwriting fonts and random "
                                       "distortions (rotation, elastic warp, noise) using tools like TextRecognitionDataGenerator; "
                                       "used for pre-training and to enlarge the training set."),
        ("Optional extra corpora", "IAM-OnDB, CVL, RIMES for additional robustness; KHATT/UNHD/UrduDoc for an "
                                   "optional Urdu/Arabic-script proof-of-concept."),
        ("Custom test set", "A small set of scanned notes/forms in the team members’ own handwriting for a "
                            "realistic end-to-end evaluation."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])
    heading(doc, "10.2  Data Pre-processing", 2, th)
    for s in [
        "Extract line images and ground-truth transcripts from the IAM XML; filter unreadable or "
        "wrongly-segmented samples; build a fixed character vocabulary (letters, digits, punctuation, space, CTC blank).",
        "Normalise each line image: grayscale, deskew, binarise (or keep grayscale), pad/resize to a fixed "
        "height (e.g. 32–64 px) while preserving aspect ratio; right-pad widths within a batch.",
        "Tokenise transcripts to integer label sequences; record sequence lengths for CTC.",
        "For multi-line documents at inference time: deskew the page, segment into lines via projection "
        "profiles/connected components, then feed each line to the model.",
    ]:
        bullet(doc, s, th)
    heading(doc, "10.3  Training / Validation / Testing Split", 2, th)
    make_table(doc, ["Split", "Proportion", "Purpose"],
        [["Training", "≈ 70–75 %", "Learn CRNN weights (IAM official train set + synthetic data)"],
         ["Validation", "≈ 10–15 %", "Hyper-parameter tuning, early stopping, decoder/LM selection"],
         ["Testing", "≈ 15 %", "Final unbiased CER/WER (IAM official test set + custom handwriting set)"]],
        th, col_widths=[1.4, 1.5, 3.5], first_col_bold=True,
        caption="Table 2.3: Data partitioning. The IAM writer-independent split is respected so that no writer "
                "appears in both training and testing.")
    heading(doc, "10.4  Data Augmentation", 2, th)
    body(doc, "Augmentation is critical because handwriting datasets are limited and writer-dependent: random "
              "rotations (±3–5°), slant/shear, elastic and grid distortions, random scaling and translation, "
              "morphological erosion/dilation (stroke thinning/thickening), brightness/contrast jitter, "
              "Gaussian/salt-pepper noise, and random background textures. Large volumes of synthetic "
              "handwriting can be generated to pre-train the network before fine-tuning on real IAM data.", th)

    section(doc, th, 11, "Expected Outcomes")
    for b in [
        "A trained CRNN+CTC recogniser with line-level CER ≈ 8–12 % on IAM without a language model, improving "
        "with beam search + SymSpell post-correction.",
        "A working document pipeline: deskew → binarise → line/word segmentation → recognition → post-correction "
        "→ layout reconstruction.",
        "Multi-language capability: an English model plus an optional Urdu/Arabic-script proof-of-concept model.",
        "Export to TXT, DOCX and searchable PDF, preserving basic line/paragraph structure.",
        "A desktop/web application with file/PDF upload, near-real-time webcam capture, and an editable side-by-side review panel.",
        "Complete documentation: code repository, training notebooks, evaluation report (CER/WER, error analysis) and user guide.",
    ]:
        bullet(doc, b, th)
    add_picture_centered(doc, CHART, 6.6, caption="Figure 2.3: (left) Character Error Rate of representative HTR "
                         "methods on the IAM test set with our expected operating point; (right) an illustrative "
                         "training/validation CER convergence curve.", th=th)

    section(doc, th, 12, "Innovation / Uniqueness")
    for b in [
        ("Complete document workflow", "Not just a recogniser — a full pipeline (pre-process → segment → "
            "recognise → correct → rebuild layout → export) wrapped in a usable app."),
        ("Hybrid accuracy boosting", "Combines a compact CRNN with CTC beam search and a dictionary/edit-"
            "distance post-corrector, getting near-Transformer usefulness at a fraction of the compute."),
        ("Multi-language design", "Pluggable model + dictionary architecture, demonstrated with English and an "
            "Urdu/Arabic-script proof-of-concept — relevant to the local context."),
        ("Low-resource & reproducible", "Lightweight CNN backbone, synthetic-data pre-training and free cloud "
            "GPUs make the whole project reproducible on student laptops."),
        ("Accessibility & preservation focus", "Editable, searchable, screen-reader-friendly output supports "
            "visually-impaired users and archival digitisation."),
        ("Practical export formats", "Direct export to DOCX and searchable PDF (OCR text layer) — features real "
            "users actually need, rarely included in coursework projects."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])

    section(doc, th, 13, "Feasibility Analysis")
    heading(doc, "13.1  Technical Feasibility", 2, th)
    body(doc, "CRNN+CTC has many open reference implementations; IAM is freely available for research; OpenCV, "
              "PyTorch, Tesseract, SymSpell, python-docx and pdf2image are mature. A compact CRNN has only a few "
              "million parameters and trains in a few hours on a free Colab T4 GPU, especially when warm-started "
              "from synthetic-data pre-training. Inference on a single text line is millisecond-scale on CPU, so "
              "the application runs comfortably on student laptops.", th)
    heading(doc, "13.2  Cost Analysis", 2, th)
    make_table(doc, ["Item", "Option Chosen", "Cost (PKR)"],
        [["Compute for training", "Google Colab / Kaggle free GPU tier", "0"],
         ["Datasets", "IAM (free for research) + synthetic data we generate", "0"],
         ["Software & libraries", "Python, PyTorch, OpenCV, Tesseract, SymSpell, python-docx (open-source)", "0"],
         ["Development machines", "Members’ existing laptops", "0 (already owned)"],
         ["Optional Colab Pro / Kaggle add-on", "Only if final tuning needs more GPU time", "≈ 2 800 / month (optional)"],
         ["Demo hosting (optional)", "Streamlit Community Cloud / Hugging Face Spaces (free)", "0"],
         ["Estimated total", "Core project", "≈ 0 (PKR 0–3 000 if Colab Pro used)"]],
        th, col_widths=[2.3, 2.6, 1.5], first_col_bold=True,
        caption="Table 2.4: Indicative cost analysis — effectively zero-budget.")
    heading(doc, "13.3  Hardware Feasibility & GPU / Resource Limitations", 2, th)
    for b in [
        "Training: a free Colab/Kaggle GPU (T4/P100) is sufficient for a compact CRNN; full TrOCR-scale "
        "Transformer training is intentionally out of scope.",
        "Local development/inference: any laptop with 8 GB+ RAM; no GPU strictly required for inference.",
        "Mitigations for limited resources: synthetic-data pre-training to converge faster, fixed small line-"
        "image height to bound memory, gradient clipping and mixed precision, and checkpoint-and-resume to "
        "survive Colab session limits.",
    ]:
        bullet(doc, b, th)
    heading(doc, "13.4  Practical Implementation Plan", 2, th)
    body(doc, "Weeks 1–4: literature, IAM download, pre-processing/segmentation pipeline, synthetic-data "
              "generator, per-character CNN baseline. Weeks 5–8: build and train the CRNN+CTC model on Colab, "
              "reach target CER, add beam-search decoding. Weeks 9–12: SymSpell post-correction, multi-language "
              "model, layout reconstruction, export module, UI. Weeks 13–16: integration, testing on custom "
              "handwriting, evaluation report, documentation and final demo. Roles: pre-processing & "
              "segmentation (member A), model & training (members A & B), post-processing & multi-language "
              "(member C), UI/export & documentation (member D), with weekly integration reviews.", th)

    section(doc, th, 14, "Hardware and Software Requirements")
    make_table(doc, ["Requirement", "Minimum", "Recommended"],
        [["Processor", "Quad-core CPU (Intel i5 / Ryzen 5)", "8-core CPU"],
         ["RAM", "8 GB", "16 GB"],
         ["GPU (local, optional)", "Integrated GPU (CPU inference fine)", "NVIDIA GPU 4 GB+ (e.g. GTX 1650)"],
         ["GPU (training)", "Google Colab / Kaggle free T4", "Colab Pro / institutional GPU"],
         ["Storage", "15 GB free", "30 GB free (IAM + synthetic data + checkpoints)"],
         ["OS", "Windows 10 / Ubuntu 20.04 / macOS", "Ubuntu 22.04"],
         ["Software", "Python 3.10, PyTorch, OpenCV, Tesseract, SymSpell, python-docx, Streamlit, Git", "+ CUDA toolkit (if local GPU), VS Code"],
         ["Peripherals", "Scanner or smartphone camera; optional webcam", "Flatbed scanner for clean page captures"]],
        th, col_widths=[1.7, 2.5, 2.2], first_col_bold=True,
        caption="Table 2.5: Hardware and software requirements.")

    section(doc, th, 15, "Future Enhancements")
    for b in [
        "Fine-tune a Transformer recogniser (TrOCR-small) for higher accuracy when more compute is available.",
        "Full-page layout analysis: multi-column text, tables, forms, marginalia, mixed print/handwriting.",
        "More languages and scripts (Sindhi, Pashto, Persian) and handwritten-mathematics recognition.",
        "On-device mobile app (TFLite/ONNX) for instant note capture and digitisation.",
        "Writer adaptation / few-shot personalisation to a specific user’s handwriting.",
        "Integration with note-taking and document-management systems; batch-folder processing.",
        "Active-learning loop where user corrections are fed back to improve the model.",
    ]:
        bullet(doc, b, th)

    section(doc, th, 16, "Timeline / Gantt Chart")
    body(doc, "The project runs over a ≈ 16-week semester in three phases — foundation, modelling & development, "
              "and finalisation — with the weekly milestones below.", th)
    add_picture_centered(doc, GANTT, 6.7, caption="Figure 2.4: Project Gantt chart with weekly milestone breakdown.", th=th)
    make_table(doc, ["Week(s)", "Milestone / Deliverable"],
        [["W1–W2", "Requirement analysis, literature review, finalise CRNN architecture and tools"],
         ["W2–W3", "Download & prepare IAM database; exploratory analysis of lines/writers"],
         ["W3–W4", "Pre-processing & line/word segmentation pipeline; synthetic-data generator"],
         ["W5–W6", "Per-character CNN baseline; set up Colab training; build CRNN skeleton"],
         ["W6–W8", "Train CRNN+CTC (pre-train on synthetic, fine-tune on IAM); reach target CER; add beam search"],
         ["W9–W10", "SymSpell post-correction; optional Urdu/Arabic-script model proof-of-concept"],
         ["W10–W12", "Layout reconstruction; export to TXT/DOCX/searchable PDF; build the UI"],
         ["W12–W13", "End-to-end integration; webcam capture mode; performance tuning"],
         ["W13–W14", "Testing on IAM + custom handwriting; CER/WER evaluation and error analysis"],
         ["W14–W15", "Documentation, user guide and final report writing"],
         ["W16", "Final review, demonstration and submission"]],
        th, col_widths=[1.1, 5.3], first_col_bold=True,
        caption="Table 2.6: Weekly milestone breakdown.")

    section(doc, th, 17, "Conclusion")
    body(doc, "This proposal presented an AI-Based Handwritten Text Recognition system that turns images, scans "
              "and PDFs of handwriting into editable, searchable text exportable as TXT, DOCX or searchable PDF. "
              "Its recognition core is the well-proven CRNN + CTC architecture, kept deliberately compact so it "
              "can be pre-trained on synthetic data and fine-tuned on the IAM database using only a free cloud "
              "GPU. The project’s value lies in delivering the complete document workflow — pre-processing, "
              "segmentation, post-correction, layout reconstruction, multi-language support and practical "
              "exports — wrapped in a usable application. It is realistic for a four-member team to complete in "
              "one semester, builds core competencies in CNNs, sequence modelling and CTC, and yields a tool "
              "with genuine utility for students, offices and archives, with a clear path toward Transformer-"
              "based and mobile future versions.", th)

    section(doc, th, 18, "References")
    refs_2 = [
        'B. Shi, X. Bai, and C. Yao, "An end-to-end trainable neural network for image-based sequence '
        'recognition and its application to scene text recognition," IEEE Trans. Pattern Analysis and Machine '
        'Intelligence, vol. 39, no. 11, pp. 2298–2304, 2017. [Online]. Available: https://arxiv.org/abs/1507.05717',
        'A. Graves, M. Liwicki, S. Fernández, R. Bertolami, H. Bunke, and J. Schmidhuber, "A novel connectionist '
        'system for unconstrained handwriting recognition," IEEE Trans. Pattern Analysis and Machine '
        'Intelligence, vol. 31, no. 5, pp. 855–868, 2009.',
        'A. Graves, S. Fernández, F. Gomez, and J. Schmidhuber, "Connectionist temporal classification: Labelling '
        'unsegmented sequence data with recurrent neural networks," in Proc. Int. Conf. Machine Learning (ICML), 2006, pp. 369–376.',
        'J. Puigcerver, "Are multidimensional recurrent layers really necessary for handwritten text '
        'recognition?," in Proc. IAPR Int. Conf. Document Analysis and Recognition (ICDAR), 2017, pp. 67–72.',
        'T. Bluche and R. Messina, "Gated convolutional recurrent neural networks for multilingual handwriting '
        'recognition," in Proc. IAPR ICDAR, 2017, pp. 646–651.',
        'M. Li, T. Lv, J. Chen, L. Cui, Y. Lu, D. Florencio, C. Zhang, Z. Li, and F. Wei, "TrOCR: Transformer-'
        'based optical character recognition with pre-trained models," in Proc. AAAI Conf. Artificial '
        'Intelligence, 2023. [Online]. Available: https://arxiv.org/abs/2109.10282',
        'U.-V. Marti and H. Bunke, "The IAM-database: An English sentence database for offline handwriting '
        'recognition," Int. Journal on Document Analysis and Recognition, vol. 5, no. 1, pp. 39–46, 2002.',
        'A. Graves, "Supervised Sequence Labelling with Recurrent Neural Networks," Studies in Computational '
        'Intelligence, vol. 385, Springer, 2012.',
        'R. Smith, "An overview of the Tesseract OCR engine," in Proc. IAPR ICDAR, 2007, pp. 629–633.',
        'A. Paszke et al., "PyTorch: An imperative style, high-performance deep learning library," in NeurIPS, 2019, pp. 8024–8035.',
        'G. Bradski, "The OpenCV Library," Dr. Dobb’s Journal of Software Tools, 2000.',
        'W. Garbe, "SymSpell: 1 million times faster spelling correction & fuzzy search through Symmetric Delete '
        'spelling correction algorithm," 2012. [Online]. Available: https://github.com/wolfgarbe/SymSpell',
        'Research Group on Computer Vision and Artificial Intelligence, University of Bern, "IAM Handwriting '
        'Database." [Online]. Available: https://fki.tic.heia-fr.ch/databases/iam-handwriting-database',
        'Y. Du et al., "PP-OCR: A practical ultra lightweight OCR system," arXiv:2009.09941, 2020.',
    ]
    _refs(doc, th, refs_2)


# ===========================================================================
# PROPOSAL 3 — SIGN LANGUAGE RECOGNITION SYSTEM
# ===========================================================================
def _proposal_3(doc, th, DIV, ARCH, FLOW, GANTT, CHART):
    section(doc, th, 2, "Abstract")
    body(doc, "Deaf and hard-of-hearing people communicate primarily through sign language, but most hearing "
              "people do not understand it, creating a daily communication barrier in schools, hospitals, banks "
              "and public offices. This project proposes a Sign Language Recognition System that uses a webcam "
              "and computer vision to translate hand-sign gestures into on-screen text and synthesised speech in "
              "real time. Rather than feeding raw images to a heavy CNN, the system uses Google’s MediaPipe Hands "
              "to extract 21 three-dimensional hand landmarks per frame; these compact, lighting- and "
              "background-robust coordinates are normalised and fed to a lightweight classifier — a small "
              "multi-layer perceptron or 1-D CNN for static signs (e.g. the ASL alphabet), and an LSTM/GRU over "
              "a short frame window for dynamic gestures (motion-based words). Recognised letters/words are "
              "assembled into sentences with an auto-complete helper and spoken aloud via an offline text-to-"
              "speech engine. Because MediaPipe is highly optimised and the classifiers are tiny, the entire "
              "pipeline runs in real time on an ordinary laptop CPU with a basic webcam — no GPU required. "
              "Targeted performance is ≥ 95 % accuracy on a held-out set of static alphabet signs and a robust, "
              "low-latency live demo, delivered as a desktop application with text and voice output.", th)
    add_picture_centered(doc, DIV, 6.5, th=th)
    body(doc, "Keywords:  Sign language recognition, MediaPipe Hands, hand landmarks, gesture recognition, "
              "LSTM, assistive technology, real-time computer vision, text-to-speech, accessibility.", th, italic=True, size=10)

    section(doc, th, 3, "Introduction")
    body(doc, "Sign languages are full natural languages expressed through hand shapes, movements, orientation, "
              "and facial expression. The World Health Organization estimates that over 400 million people "
              "worldwide have disabling hearing loss; in Pakistan, Pakistan Sign Language (PSL) serves a large "
              "Deaf community. Yet outside specialised settings there are very few interpreters, so routine "
              "interactions — a doctor’s appointment, a bank transaction, a classroom question — often break "
              "down. An automatic system that recognises signs from a camera and renders them as text and speech "
              "can act as a always-available, low-cost interpreter for many everyday situations.", th)
    body(doc, "Computer-vision approaches to sign recognition have ranged from glove-based sensors and depth "
              "cameras (Kinect) to deep CNNs on raw RGB images and 3-D CNNs / I3D on video clips. Image-based "
              "CNNs are sensitive to lighting, skin tone, clothing and background, and video models are "
              "compute-heavy. A more practical route — and the one this project takes — is landmark-based "
              "recognition: a pretrained hand-tracking model (MediaPipe Hands, based on the BlazePalm/hand-"
              "landmark networks of Zhang et al., 2020) outputs 21 normalised 3-D keypoints per detected hand at "
              "high frame rates on CPU. Classifying these low-dimensional, invariant features needs only a tiny "
              "neural network, which is fast to train, fast to run, and far more robust to nuisance variation. "
              "Dynamic, motion-based signs are handled by adding a short temporal model (LSTM/GRU) over a sliding "
              "window of landmark frames.", th)
    body(doc, "The scope is sized for a one-semester, four-member project: focus on the ASL/PSL fingerspelling "
              "alphabet plus a curated set of common dynamic words, build a clean real-time UI, and add sentence "
              "construction and offline speech output. The emphasis is on a robust, genuinely usable assistive "
              "prototype rather than maximising vocabulary size.", th)

    section(doc, th, 4, "Problem Statement")
    body(doc, "Communication between Deaf signers and the hearing majority is obstructed by a chronic shortage "
              "of interpreters and the general public’s lack of sign-language knowledge. Existing technical aids "
              "have drawbacks:", th)
    bullet(doc, "Glove/sensor solutions are expensive, intrusive and not always available.", th, bold_lead="Hardware-heavy")
    bullet(doc, "Raw-image CNN recognisers are brittle to lighting, background, skin tone and camera quality, "
                "and often need a GPU.", th, bold_lead="Fragile / costly")
    bullet(doc, "Video (3-D CNN) models are accurate but too compute-intensive to run live on a typical laptop.", th, bold_lead="Slow")
    bullet(doc, "Most demos recognise isolated letters only and stop there — they don’t build words/sentences "
                "or produce speech.", th, bold_lead="Incomplete")
    body(doc, "Formally: given a live webcam stream, detect the signing hand(s), recognise the current static "
              "sign (e.g. an alphabet letter) or dynamic gesture (a motion-based word), assemble the recognised "
              "tokens into readable text with basic word completion, and output that text both on screen and as "
              "synthesised speech — all in real time on commodity, GPU-free hardware. This project delivers that "
              "end-to-end assistive pipeline.", th)

    section(doc, th, 5, "Objectives")
    for o in [
        ("Primary", "Build a real-time webcam pipeline that detects hands (MediaPipe), recognises static "
                    "ASL/PSL alphabet signs with ≥ 95 % accuracy on a held-out test set, and runs at ≥ 15 FPS "
                    "on a laptop CPU."),
        ("Secondary", "Add a dynamic-gesture module (LSTM/GRU over a short landmark-frame window) for a curated "
                      "set of common motion-based signs/words."),
        ("Secondary", "Implement landmark normalisation (translation/scale invariance, optional handedness "
                      "normalisation) so recognition is robust to position, distance and background."),
        ("Secondary", "Build a sentence constructor with debouncing/temporal smoothing and an auto-complete/"
                      "word-suggestion helper to turn recognised letters into fluent text."),
        ("Secondary", "Integrate offline text-to-speech (pyttsx3) — with optional online gTTS — to speak the "
                      "constructed text aloud."),
        ("Secondary", "Deliver a clean desktop UI showing the webcam feed with hand-landmark overlay, the "
                      "current prediction with confidence, the building sentence, and play/clear controls."),
        ("Learning", "Demonstrate competence in real-time computer vision, feature engineering from landmarks, "
                     "lightweight classifiers, sequence models (LSTM/GRU), and assistive-application design."),
    ]:
        bullet(doc, o[1], th, bold_lead=o[0] + " objective")

    section(doc, th, 6, "Scope of the Project")
    heading(doc, "In Scope", 2, th)
    for s in [
        "Real-time recognition of the static ASL/PSL fingerspelling alphabet (A–Z, allowing for J/Z motion) from a webcam.",
        "A dynamic-gesture module for a curated set (≈ 10–30) of common motion-based words (e.g. hello, thanks, yes, no, please).",
        "MediaPipe-based hand detection and 21-landmark extraction; landmark normalisation and feature engineering.",
        "Lightweight classifiers (MLP / 1-D CNN for static; LSTM/GRU for dynamic) trained on public + self-collected data.",
        "Sentence construction with temporal smoothing/debouncing and an auto-complete helper.",
        "On-screen text output plus offline text-to-speech; a desktop application UI with landmark overlay and controls.",
    ]:
        bullet(doc, s, th)
    heading(doc, "Out of Scope (Limitations)", 2, th)
    for s in [
        "Full continuous sign-language translation with grammar, facial expressions and body pose (a much larger research problem).",
        "Large-vocabulary word recognition (e.g. full WLASL-2000) — only a curated subset is targeted.",
        "Two-hand complex co-articulated signs beyond the curated set; signer-independent perfection in all conditions.",
        "Mobile/embedded deployment (feasible later) and speech-to-sign (the reverse direction).",
        "Production-grade robustness in extreme lighting, motion blur or heavy occlusion.",
    ]:
        bullet(doc, s, th)

    section(doc, th, 7, "Literature Review")
    body(doc, "Sign language recognition (SLR) research spans sensor-based methods, image-based deep learning, "
              "video-based temporal models, and — increasingly — skeleton/landmark-based approaches that mirror "
              "trends in human-action recognition. The works below justify our landmark-centric, lightweight "
              "design.", th)
    heading(doc, "7.1  Key Works", 2, th)
    body(doc, "Early SLR used data gloves and accelerometers, and later depth cameras (Microsoft Kinect) to "
              "obtain hand/skeleton data; these are accurate but require special hardware. With deep learning, "
              "CNNs were applied directly to RGB hand images (e.g. ASL-alphabet classifiers), and for dynamic "
              "signs, 3-D CNNs and Inflated 3D ConvNets (I3D; Carreira & Zisserman, 2017) and CNN-LSTM hybrids "
              "modelled motion in video; Li et al. (2020) released WLASL, a large word-level ASL video benchmark, "
              "with I3D baselines. In parallel, skeleton-based action recognition (e.g. ST-GCN, Yan et al., 2018) "
              "showed that joint coordinates alone carry strong discriminative signal. The enabling technology "
              "for our approach is MediaPipe Hands (Zhang et al., 2020): a real-time, on-device hand-tracking "
              "pipeline (palm detector + hand-landmark regressor) producing 21 3-D landmarks per hand at high "
              "frame rates on CPU/mobile. Building classifiers on these landmarks — MLPs and 1-D CNNs for static "
              "poses, LSTM/GRU or lightweight Transformers/GCNs for dynamic sequences — is now a common, "
              "efficient SLR recipe; numerous IEEE/Elsevier papers report 95–99 % accuracy on alphabet datasets "
              "with such pipelines. Text-to-speech is handled by standard engines (pyttsx3 offline; Google TTS "
              "online).", th)
    heading(doc, "7.2  Comparison of SLR Approaches", 2, th)
    make_table(doc,
        ["Approach", "Input / Features", "Strength", "Limitation / Cost", "Typical Accuracy*"],
        [
            ["Sensor glove / accelerometer", "Bend/flex & motion sensors", "Very robust signal",
             "Special hardware; intrusive; costly", "high (but impractical)"],
            ["Depth camera (Kinect) + ML", "Depth + skeleton joints", "3-D info; lighting-robust",
             "Requires depth sensor", "≈ 90–96 %"],
            ["CNN on raw RGB hand image", "Cropped hand image", "Simple; no landmark library",
             "Brittle to light/background/skin; often needs GPU", "≈ 90–95 %"],
            ["3-D CNN / I3D on video", "Stacked RGB frames", "Captures motion well (dynamic signs)",
             "Compute-heavy; not real-time on CPU", "≈ 90–96 % (WLASL subsets lower)"],
            ["MediaPipe landmarks + MLP/1-D CNN ★", "21×3 normalised keypoints (static)", "Tiny, fast, robust, CPU real-time",
             "Static poses only; depends on landmark quality", "≈ 96–99 % (alphabet)"],
            ["MediaPipe landmarks + LSTM/GRU ★", "Landmark sequence over a window (dynamic)", "Handles motion; still lightweight",
             "Needs sequence data collection", "≈ 95–98 % (curated words)"],
            ["Skeleton GCN (ST-GCN-style)", "Graph over hand/body joints", "Strong for complex gestures",
             "More complex; more data", "≈ 95–98 %"],
        ], th, col_widths=[1.55, 1.5, 1.45, 1.55, 1.05], header_fs=8.6, body_fs=8.5,
        caption="Table 3.1: SLR approaches. ★ = methods adopted (MediaPipe landmarks + lightweight static and "
                "dynamic classifiers). *Accuracies are indicative ranges reported across alphabet/curated-word datasets; "
                "they are not directly comparable across different datasets.")
    heading(doc, "7.3  Datasets & Resources", 2, th)
    make_table(doc, ["Resource", "Content", "Use Here"],
        [["ASL Alphabet (Kaggle)", "≈ 87 000 images, 29 classes (A–Z + space/del/nothing)", "Train/test the static-sign classifier"],
         ["Sign Language MNIST", "≈ 34 600 28×28 images, 24 static letters", "Quick prototyping baseline"],
         ["WLASL (subset)", "Word-level ASL videos (2 000 words; we use a small subset)", "Optional dynamic-gesture data"],
         ["Self-collected landmark data", "Team members record alphabet + curated words via webcam → save MediaPipe landmarks",
          "Domain-matched data; covers PSL signs not in public sets"],
         ["MediaPipe Hands, pyttsx3/gTTS, NLTK", "Pretrained hand tracker; TTS engines; word lists", "Feature extraction, speech output, auto-complete"]],
        th, col_widths=[1.9, 2.5, 2.0], header_fs=8.8, body_fs=8.6,
        caption="Table 3.2: Datasets and resources considered.")
    body(doc, "Research gap addressed:  many projects classify isolated ASL letters from images and stop. Fewer "
              "deliver a robust landmark-based pipeline that also handles dynamic gestures, builds sentences with "
              "smoothing and auto-complete, speaks the output, and runs in real time on a GPU-free laptop — and "
              "fewer still consider Pakistan Sign Language. This project targets that practical, accessibility-"
              "focused gap.", th, italic=True)

    section(doc, th, 8, "Proposed Methodology")
    body(doc, "The methodology comprises a data-collection/feature-engineering stage, a model-training stage "
              "(separate static and dynamic models), and a real-time application/integration stage.", th)
    heading(doc, "8.1  System Architecture", 2, th)
    add_picture_centered(doc, ARCH, 6.6, caption="Figure 3.1: End-to-end architecture — webcam → MediaPipe "
                         "21-landmark extraction → normalisation → static MLP/1-D CNN and dynamic LSTM/GRU "
                         "classifiers → sentence builder → on-screen text + text-to-speech.", th=th)
    heading(doc, "8.2  Step-by-Step Workflow", 2, th)
    for lead, txt in [
        ("Data collection", "Use the ASL Alphabet/Sign-Language-MNIST datasets for static signs and record "
            "additional webcam clips (team + volunteers) for PSL signs and dynamic words; for every frame, run "
            "MediaPipe Hands and store the 21×3 landmark array (plus handedness) with the label."),
        ("Landmark normalisation & features", "Translate landmarks so the wrist is the origin; scale by a "
            "reference distance (e.g. wrist-to-middle-finger-MCP) for distance invariance; optionally mirror "
            "left hands to a canonical handedness; derive extra features (inter-joint distances/angles). For "
            "dynamic signs, stack normalised landmarks over a sliding window (e.g. 20–30 frames)."),
        ("Static classifier", "Train a small MLP (or 1-D CNN over the 63-value vector) to classify alphabet "
            "signs; use dropout and class balancing; this network has only a few thousand–tens-of-thousands of "
            "parameters and trains in minutes on CPU."),
        ("Dynamic classifier", "Train an LSTM/GRU (1–2 layers) over the landmark-sequence window to classify "
            "the curated motion-based words; pad/truncate sequences; apply temporal data augmentation."),
        ("Real-time inference", "Capture webcam frames with OpenCV; run MediaPipe; normalise; route to the "
            "static model each frame and to the dynamic model on the rolling window; obtain class + confidence."),
        ("Temporal smoothing / debouncing", "Stabilise output with a majority vote / confidence threshold over "
            "the last k frames so a letter is only ‘committed’ when held steady; suppress jitter and the "
            "no-hand state."),
        ("Sentence construction", "Append committed letters/words to a text buffer; map special signs to space/"
            "backspace; offer auto-complete word suggestions (frequency dictionary / prefix trie)."),
        ("Text-to-speech", "On demand (or sentence-end), speak the buffer aloud using pyttsx3 (offline) or gTTS "
            "(online); show speaking status in the UI."),
        ("UI integration", "Display the webcam feed with the hand-landmark skeleton overlay, the current "
            "prediction and confidence bar, the building sentence, FPS, and buttons for speak / clear / "
            "static-vs-dynamic mode."),
        ("Testing & evaluation", "Measure per-class accuracy and a confusion matrix on held-out data, plus a "
            "live user test (recognition rate, latency, FPS) with several signers; document results."),
    ]:
        numbered(doc, txt, th, bold_lead=lead)
    heading(doc, "8.3  Workflow Flowchart", 2, th)
    add_picture_centered(doc, FLOW, 4.3, caption="Figure 3.2: Project workflow with the accuracy decision gate "
                         "and refinement loop.", th=th)
    heading(doc, "8.4  Algorithms, Models & Frameworks", 2, th)
    for b in [
        ("MediaPipe Hands (pretrained)", "Palm detector + hand-landmark model → 21 3-D landmarks per hand, real-time on CPU."),
        ("Static classifier", "Small MLP (e.g. 63→128→64→#classes) or 1-D CNN over the landmark vector."),
        ("Dynamic classifier", "1–2-layer LSTM/GRU over a window of normalised landmark frames (optional lightweight Transformer/GCN)."),
        ("Smoothing", "Sliding-window majority vote + confidence thresholding (debouncing)."),
        ("Text-to-speech", "pyttsx3 (offline, default) or gTTS + playback (online option)."),
        ("Auto-complete", "Frequency-ranked dictionary / prefix trie for word suggestions."),
        ("Frameworks", "Python, OpenCV, MediaPipe, NumPy, scikit-learn, TensorFlow-Keras or PyTorch (for LSTM), pyttsx3/gTTS, Tkinter/PyQt or Streamlit."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])

    section(doc, th, 9, "Technologies and Tools")
    _common_tech_table(doc, th, [
        ["Programming language", "Python 3.10+", "Vision pipeline, model training, UI and integration"],
        ["Hand tracking", "Google MediaPipe (Hands solution)", "Pretrained real-time 21-landmark extraction on CPU"],
        ["Computer vision", "OpenCV", "Webcam capture, frame handling, drawing landmark overlays"],
        ["ML / DL frameworks", "scikit-learn; TensorFlow-Keras or PyTorch", "Static MLP/1-D CNN and dynamic LSTM/GRU models"],
        ["Numerical", "NumPy, pandas", "Landmark arrays, feature engineering, dataset handling"],
        ["Text-to-speech", "pyttsx3 (offline); gTTS + playsound (online option)", "Speaking the constructed sentences aloud"],
        ["NLP helper", "NLTK / wordfreq, prefix trie", "Auto-complete and word suggestions"],
        ["Application UI", "Tkinter / PyQt (alt. Streamlit)", "Live feed, prediction display, sentence panel, controls"],
        ["Visualisation", "Matplotlib, seaborn", "Confusion matrices, accuracy curves, evaluation plots"],
        ["Training environment", "Local CPU; Google Colab (free) if needed", "Tiny models train on CPU; Colab optional for LSTM"],
        ["Dev tools", "VS Code, Jupyter, Git/GitHub", "Coding, experiments, version control, collaboration"],
        ["Dataset sources", "ASL Alphabet & Sign-Language-MNIST (Kaggle); WLASL subset; self-collected landmarks",
         "Static training/testing; optional dynamic data; domain-matched PSL data"],
    ])

    section(doc, th, 10, "Dataset Information")
    heading(doc, "10.1  Candidate Datasets", 2, th)
    for b in [
        ("ASL Alphabet (Kaggle)", "≈ 87 000 RGB images across 29 classes (A–Z plus space/delete/nothing); the "
                                  "main source for the static-sign classifier (we process images through MediaPipe to obtain landmarks)."),
        ("Sign Language MNIST", "≈ 34 600 grayscale 28×28 images of 24 static letters; a quick baseline/sanity-check dataset."),
        ("WLASL (subset)", "Word-level ASL video clips; a small curated subset (common words) for the dynamic-gesture module."),
        ("Self-collected landmark dataset", "Team members and volunteers record the alphabet and ≈ 10–30 dynamic "
                                            "words via webcam; for each frame MediaPipe landmarks (21×3 + handedness) are saved with the label — "
                                            "this covers PSL signs and matches our deployment conditions."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])
    heading(doc, "10.2  Data Pre-processing", 2, th)
    for s in [
        "Run MediaPipe Hands on every image/frame; discard frames with no confident hand detection; keep the "
        "21×3 landmark array (and handedness) as the feature record.",
        "Normalise: subtract the wrist coordinate (translation invariance); divide by a reference bone length "
        "(scale invariance); optionally mirror left hands; flatten to a 63-value vector (static) or stack into "
        "a [T×63] tensor (dynamic).",
        "Engineer optional extra features: pairwise fingertip distances and joint angles to help the classifier.",
        "Balance classes; shuffle; for dynamic data, segment continuous recordings into fixed-length windows "
        "with overlap; encode labels.",
    ]:
        bullet(doc, s, th)
    heading(doc, "10.3  Training / Validation / Testing Split", 2, th)
    make_table(doc, ["Split", "Proportion", "Purpose"],
        [["Training", "≈ 70 %", "Fit the static MLP/1-D CNN and the dynamic LSTM/GRU"],
         ["Validation", "≈ 15 %", "Hyper-parameter tuning, early stopping, model selection"],
         ["Testing", "≈ 15 %", "Final per-class accuracy and confusion matrix; held-out signers where possible"]],
        th, col_widths=[1.4, 1.4, 3.6], first_col_bold=True,
        caption="Table 3.3: Data partitioning. Where enough signers are available, the split is signer-"
                "independent so that no signer appears in both training and testing.")
    heading(doc, "10.4  Data Augmentation", 2, th)
    body(doc, "Working in landmark space makes augmentation cheap and effective: small random rotations of the "
              "landmark cloud (in-plane and slight 3-D tilt), random scaling and translation, mild Gaussian "
              "jitter on each coordinate (sensor-noise simulation), random horizontal mirroring (handedness), "
              "and — for dynamic gestures — temporal jitter (frame dropping/duplication, speed warping, window "
              "shifting). This multiplies the effective dataset size and improves robustness to camera angle, "
              "distance and signing style without collecting more raw video.", th)

    section(doc, th, 11, "Expected Outcomes")
    for b in [
        "A real-time webcam application that overlays hand landmarks and recognises static ASL/PSL alphabet "
        "signs with ≥ 95 % held-out accuracy at ≥ 15 FPS on a laptop CPU.",
        "A dynamic-gesture module recognising a curated set of common motion-based words via an LSTM/GRU over landmark sequences.",
        "Stable output thanks to temporal smoothing/debouncing, with the recognised letters/words assembled into sentences.",
        "An auto-complete/word-suggestion helper that speeds up sentence construction.",
        "Offline text-to-speech (pyttsx3) that speaks the constructed text aloud; optional online gTTS.",
        "A clean desktop UI (live feed + landmark overlay, prediction + confidence, sentence panel, FPS, controls).",
        "Complete documentation: code repository, data-collection scripts, evaluation report (accuracy, confusion matrix, latency/FPS) and user guide.",
    ]:
        bullet(doc, b, th)
    add_picture_centered(doc, CHART, 6.6, caption="Figure 3.3: (left) Accuracy of representative sign-recognition "
                         "approaches with our landmark-based targets; (right) an illustrative confusion matrix "
                         "for a subset of alphabet classes.", th=th)

    section(doc, th, 12, "Innovation / Uniqueness")
    for b in [
        ("Landmark-first, GPU-free design", "Uses pretrained MediaPipe landmarks + tiny classifiers so the whole "
            "system runs in real time on a plain laptop CPU — robust to lighting/background/skin tone."),
        ("Static + dynamic in one system", "Combines an instantaneous static-sign classifier with an LSTM/GRU "
            "for motion-based signs, instead of handling only isolated letters."),
        ("Full assistive pipeline", "Goes beyond classification to sentence building (with smoothing and auto-"
            "complete) and spoken output — a usable communication aid, not just a demo."),
        ("Pakistan-Sign-Language awareness", "Self-collected data includes PSL signs, addressing the local "
            "community and a gap in public datasets."),
        ("Reproducible & low-cost", "Free datasets and tools, tiny models trainable on CPU, optional Colab — "
            "fully reproducible by students with zero budget."),
        ("Robustness engineering", "Landmark normalisation (translation/scale/handedness invariance) plus "
            "landmark-space augmentation make recognition resilient to camera angle, distance and signer style."),
    ]:
        bullet(doc, b[1], th, bold_lead=b[0])

    section(doc, th, 13, "Feasibility Analysis")
    heading(doc, "13.1  Technical Feasibility", 2, th)
    body(doc, "MediaPipe Hands is a free, well-documented, highly optimised library that already solves the "
              "hardest part (robust real-time hand landmarking). The classifiers on top are small enough to "
              "train on a CPU in minutes (static) or with a short Colab session (dynamic LSTM). OpenCV, "
              "scikit-learn/Keras, pyttsx3 and Tkinter/PyQt are all mature. Numerous published projects confirm "
              "95–99 % alphabet accuracy with this exact recipe, so the targets are realistic.", th)
    heading(doc, "13.2  Cost Analysis", 2, th)
    make_table(doc, ["Item", "Option Chosen", "Cost (PKR)"],
        [["Compute", "Local laptop CPU (training & inference); Colab free if needed", "0"],
         ["Datasets", "ASL Alphabet / Sign-Language-MNIST / WLASL subset (public) + self-collected", "0"],
         ["Software & libraries", "Python, MediaPipe, OpenCV, scikit-learn/Keras, pyttsx3 (open-source)", "0"],
         ["Camera", "Built-in laptop webcam (or an inexpensive USB webcam)", "0–2 500 (optional)"],
         ["Development machines", "Members’ existing laptops", "0 (already owned)"],
         ["Demo packaging (optional)", "PyInstaller executable / Streamlit Cloud (free)", "0"],
         ["Estimated total", "Core project", "≈ 0 (PKR 0–2 500 if a USB webcam is bought)"]],
        th, col_widths=[2.3, 2.6, 1.5], first_col_bold=True,
        caption="Table 3.4: Indicative cost analysis — essentially zero-budget.")
    heading(doc, "13.3  Hardware Feasibility & GPU / Resource Limitations", 2, th)
    for b in [
        "No GPU is required: MediaPipe runs in real time on CPU and the classifiers are tiny; a Colab GPU is "
        "only an optional convenience for training the dynamic LSTM.",
        "Any laptop with a webcam, 4–8 GB RAM and a modern dual/quad-core CPU is sufficient; FPS can be tuned by "
        "limiting MediaPipe to one hand and a moderate resolution.",
        "Resource mitigations: cache extracted landmarks (don’t re-run MediaPipe during training), keep models "
        "small, use early stopping, and design the UI loop to skip frames gracefully under load.",
    ]:
        bullet(doc, b, th)
    heading(doc, "13.4  Practical Implementation Plan", 2, th)
    body(doc, "Weeks 1–4: literature, dataset download, MediaPipe pipeline, landmark extraction & normalisation, "
              "baseline static classifier. Weeks 5–8: collect custom data, train/tune the static MLP/1-D CNN to "
              "target accuracy, build the dynamic-data recorder. Weeks 9–12: train the dynamic LSTM/GRU, add "
              "temporal smoothing, sentence builder, auto-complete and text-to-speech. Weeks 13–16: real-time UI "
              "integration, optimisation, user testing with several signers, evaluation report, documentation "
              "and final demo. Roles: MediaPipe pipeline & features (member A), static & dynamic models "
              "(members A & B), sentence builder/TTS/auto-complete (member C), UI integration & documentation "
              "(member D), with weekly integration reviews.", th)

    section(doc, th, 14, "Hardware and Software Requirements")
    make_table(doc, ["Requirement", "Minimum", "Recommended"],
        [["Processor", "Dual/Quad-core CPU (Intel i3/i5 / Ryzen 3/5)", "Quad-core+ CPU"],
         ["RAM", "4 GB", "8 GB+"],
         ["GPU", "Not required", "Optional NVIDIA GPU only for faster LSTM training"],
         ["Camera", "Built-in laptop webcam (≥ 480p)", "720p USB webcam"],
         ["Storage", "5 GB free", "15 GB free (datasets + cached landmarks + models)"],
         ["OS", "Windows 10 / Ubuntu 20.04 / macOS", "Ubuntu 22.04 / Windows 11"],
         ["Software", "Python 3.10, MediaPipe, OpenCV, scikit-learn/Keras, pyttsx3, Tkinter/PyQt, Git", "+ Jupyter, VS Code, Colab account"],
         ["Audio", "Working speakers/headphones (for TTS output)", "—"]],
        th, col_widths=[1.6, 2.6, 2.2], first_col_bold=True,
        caption="Table 3.5: Hardware and software requirements.")

    section(doc, th, 15, "Future Enhancements")
    for b in [
        "Expand the dynamic vocabulary toward full Pakistan Sign Language phrases; build a public PSL landmark dataset.",
        "Add facial-expression and body-pose cues (MediaPipe Holistic) for grammatically richer recognition.",
        "Move toward continuous sentence-level translation (sequence-to-sequence with a language model).",
        "Mobile app (MediaPipe + TFLite) for an always-available pocket interpreter.",
        "Two-way communication: speech-to-text plus animated/avatar sign output for the hearing-to-Deaf direction.",
        "Signer adaptation / few-shot personalisation; on-device continual learning from user corrections.",
        "Integration with video-call platforms for live captioning of signers.",
    ]:
        bullet(doc, b, th)

    section(doc, th, 16, "Timeline / Gantt Chart")
    body(doc, "The project runs over a ≈ 16-week semester in three phases — foundation, modelling & development, "
              "and finalisation — with the weekly milestones below.", th)
    add_picture_centered(doc, GANTT, 6.7, caption="Figure 3.4: Project Gantt chart with weekly milestone breakdown.", th=th)
    make_table(doc, ["Week(s)", "Milestone / Deliverable"],
        [["W1–W2", "Requirement analysis, literature review, finalise architecture and tools"],
         ["W2–W3", "Download ASL/Sign-MNIST data; build MediaPipe landmark-extraction pipeline; EDA"],
         ["W3–W4", "Landmark normalisation & feature engineering; data-collection scripts; baseline static classifier"],
         ["W5–W6", "Collect custom alphabet/PSL data; train & tune static MLP/1-D CNN to target accuracy"],
         ["W6–W8", "Build dynamic-gesture recorder; assemble windowed sequence dataset; train LSTM/GRU"],
         ["W9–W10", "Temporal smoothing/debouncing; sentence builder with auto-complete"],
         ["W10–W12", "Integrate text-to-speech; build the real-time desktop UI (overlay, prediction, controls)"],
         ["W12–W13", "End-to-end integration; real-time optimisation (FPS, latency)"],
         ["W13–W14", "Testing with multiple signers; accuracy/confusion-matrix/latency evaluation"],
         ["W14–W15", "Documentation, user guide and final report writing"],
         ["W16", "Final review, demonstration and submission"]],
        th, col_widths=[1.1, 5.3], first_col_bold=True,
        caption="Table 3.6: Weekly milestone breakdown.")

    section(doc, th, 17, "Conclusion")
    body(doc, "This proposal presented a Sign Language Recognition System that translates webcam hand-sign "
              "gestures into on-screen text and spoken output in real time. By building on Google’s pretrained "
              "MediaPipe Hands to obtain compact, robust 21-point landmarks and classifying them with tiny "
              "neural networks — an MLP/1-D CNN for static signs and an LSTM/GRU for dynamic ones — the system "
              "achieves high accuracy while running comfortably on a GPU-free laptop. Beyond classification it "
              "delivers a complete assistive pipeline: temporal smoothing, sentence construction with auto-"
              "complete, and offline text-to-speech, with optional support for Pakistan Sign Language. The "
              "project is realistic for a four-member team within one semester, builds strong skills in real-"
              "time computer vision and sequence modelling, and produces a socially valuable communication aid "
              "with a clear roadmap toward continuous translation and mobile deployment.", th)

    section(doc, th, 18, "References")
    refs_3 = [
        'F. Zhang, V. Bazarevsky, A. Vakunov, A. Tkachenka, G. Sung, C.-L. Chang, and M. Grundmann, "MediaPipe '
        'Hands: On-device real-time hand tracking," in CVPR Workshop on Computer Vision for Augmented and '
        'Virtual Reality, 2020. [Online]. Available: https://arxiv.org/abs/2006.10214',
        'C. Lugaresi et al., "MediaPipe: A framework for building perception pipelines," arXiv:1906.08172, 2019.',
        'D. Li, C. Rodriguez, X. Yu, and H. Li, "Word-level deep sign language recognition from video: A new '
        'large-scale dataset and methods comparison (WLASL)," in Proc. IEEE Winter Conf. on Applications of '
        'Computer Vision (WACV), 2020, pp. 1459–1469.',
        'J. Carreira and A. Zisserman, "Quo vadis, action recognition? A new model and the Kinetics dataset '
        '(I3D)," in Proc. IEEE/CVF CVPR, 2017, pp. 6299–6308.',
        'S. Yan, Y. Xiong, and D. Lin, "Spatial temporal graph convolutional networks for skeleton-based action '
        'recognition (ST-GCN)," in Proc. AAAI Conf. on Artificial Intelligence, 2018.',
        'O. Koller, J. Forster, and H. Ney, "Continuous sign language recognition: Towards large vocabulary '
        'statistical recognition systems handling multiple signers," Computer Vision and Image Understanding, '
        'vol. 141, pp. 108–125, 2015.',
        'N. C. Camgoz, S. Hadfield, O. Koller, H. Ney, and R. Bowden, "Neural sign language translation," in '
        'Proc. IEEE/CVF CVPR, 2018, pp. 7784–7793.',
        'S. Hochreiter and J. Schmidhuber, "Long short-term memory," Neural Computation, vol. 9, no. 8, '
        'pp. 1735–1780, 1997.',
        'IEEE Xplore, "Sign language recognition using deep learning — search results." [Online]. Available: '
        'https://ieeexplore.ieee.org/search/searchresult.jsp?queryText=sign+language+recognition+deep+learning',
        'G. Bradski, "The OpenCV Library," Dr. Dobb’s Journal of Software Tools, 2000.',
        'F. Chollet et al., "Keras," 2015. [Online]. Available: https://keras.io',
        'F. Pedregosa et al., "Scikit-learn: Machine learning in Python," Journal of Machine Learning Research, '
        'vol. 12, pp. 2825–2830, 2011.',
        'N. M. Bhagat (Akash) and contributors, "pyttsx3 — offline text-to-speech conversion library for '
        'Python," documentation, 2023. [Online]. Available: https://pyttsx3.readthedocs.io',
        'Akash / Kaggle, "ASL Alphabet dataset." [Online]. Available: https://www.kaggle.com/datasets/grassknoted/asl-alphabet',
    ]
    _refs(doc, th, refs_3)


# ---------------------------------------------------------------------------
def _refs(doc, th, refs):
    body(doc, "References are listed in IEEE citation style; arXiv / dataset URLs are provided where applicable.",
         th, italic=True, size=9.5, space_after=6)
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.42)
        p.paragraph_format.first_line_indent = Inches(-0.42)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.1
        rn = p.add_run(f"[{i}]\t"); set_run(rn, size=9.5, bold=True, color=th["dark"])
        rr = p.add_run(r); set_run(rr, size=9.5, color="222222")

# ===========================================================================
# Assembly
# ===========================================================================
def build_individual(pid):
    th = THEMES[pid]
    doc = Document()
    setup_document(doc)
    add_header_footer(doc.sections[0], th, f"Project Proposal {pid}: {th['short']}")
    cover_page(doc, th, pid, master=False,
               subtitle="A Final-Year Machine-Learning Project Proposal")
    table_of_contents(doc, th)
    proposal_body(doc, pid)
    enable_update_fields(doc)
    out = os.path.join(OUT, f"Proposal_{pid}.docx")
    doc.save(out)
    print("saved", out)
    return out

def build_master():
    doc = Document()
    setup_document(doc)
    th0 = THEMES[1]
    add_header_footer(doc.sections[0], th0, "Compiled Project Proposals — RISE MSAI")
    # Master cover (use proposal-1 banner but generic title)
    cover_page(doc, th0, 1, master=True,
               subtitle="Compiled Machine-Learning Project Proposals for Final-Year Selection")
    # Master TOC
    table_of_contents(doc, th0, title="Master Table of Contents", levels="1-2")
    # Overview page
    heading(doc, "About This Document", 1, th0, number="0")
    body(doc, "This master document compiles three candidate AI/ML project proposals prepared by the group for "
              "the Machine Learning course at Riphah International University (RISE MSAI). Each proposal is "
              "self-contained and follows the same academic structure — cover page, abstract, introduction, "
              "problem statement, objectives, scope, literature review, methodology, technologies, dataset "
              "information, expected outcomes, innovation, feasibility, requirements, future enhancements, "
              "timeline, conclusion and references. All three projects are deliberately scoped to be "
              "implementable by a four-member student team within one semester, using lightweight models, "
              "transfer learning, pretrained components and free cloud GPUs — without high-end hardware.", th0)
    make_table(doc, ["#", "Project Title", "Core Technique", "Key Datasets", "Headline Target"],
        [["1", "AI-Based Crowd Counting System", "CSRNet density-map regression (transfer learning, frozen VGG-16)",
          "ShanghaiTech, Mall", "MAE ≈ 70–95 (SH-A); real-time dashboard"],
         ["2", "AI-Based Handwritten Text Recognition", "CRNN (CNN + BiLSTM + CTC) + SymSpell post-correction",
          "IAM + synthetic handwriting", "CER ≈ 8–12 % on IAM; TXT/DOCX/PDF export"],
         ["3", "Sign Language Recognition System", "MediaPipe 21-landmarks + MLP/1-D CNN (static) & LSTM/GRU (dynamic)",
          "ASL Alphabet, Sign-MNIST, WLASL subset, self-collected", "≥ 95 % alphabet accuracy; real-time CPU + TTS"]],
        th0, col_widths=[0.35, 1.85, 1.95, 1.55, 1.8], header_fs=8.8, body_fs=8.6,
        caption="Table 0.1: Summary of the three proposed projects in this document.")
    body(doc, "The detailed proposals follow, beginning on the next page.", th0, italic=True)
    page_break(doc)
    # Each proposal as its own section so headers/footers can carry its theme & title
    for pid in (1, 2, 3):
        th = THEMES[pid]
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        # reset margins/page for the new section
        sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.4); sec.right_margin = Cm(2.4)
        sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
        sec.header_distance = Cm(1.1); sec.footer_distance = Cm(1.1)
        add_header_footer(sec, th, f"Proposal {pid}: {th['short']}", different_first=True)
        # internal cover page for this proposal
        cover_page(doc, th, pid, master=False,
                   subtitle=f"Proposal {pid} of 3  —  A Final-Year Machine-Learning Project Proposal")
        # mini per-proposal contents
        heading(doc, f"Contents — Proposal {pid}: {th['name']}", 1, th, number=str(pid) + ".0")
        body(doc, "This proposal contains the standard sections (Abstract through References) as listed in the "
                  "Master Table of Contents at the front of this document.", th, italic=True)
        proposal_body(doc, pid)
    enable_update_fields(doc)
    out = os.path.join(OUT, "Master_Proposal_Document.docx")
    doc.save(out)
    print("saved", out)
    return out

def main():
    for pid in (1, 2, 3):
        build_individual(pid)
    build_master()

if __name__ == "__main__":
    main()
